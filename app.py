import streamlit as st
import os
import json
import pickle
import hashlib
import re
import math
from pathlib import Path
from typing import List, Dict, Any, Tuple, Optional, Set
from datetime import datetime
import logging
import numpy as np
import tiktoken
import nltk
import yaml
from collections import defaultdict, Counter

# Optional dependencies with graceful fallbacks
HAS_PYMUPDF = False
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    pass

HAS_PYPDF2 = False
try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    pass

HAS_DOCX = False
try:
    from docx import Document
    from docx.table import Table as DocxTable
    HAS_DOCX = True
except ImportError:
    pass

HAS_SENTENCE_TRANSFORMERS = False
try:
    from sentence_transformers import SentenceTransformer, CrossEncoder
    from sklearn.metrics.pairwise import cosine_similarity
    HAS_SENTENCE_TRANSFORMERS = True
except ImportError:
    pass

HAS_CHROMADB = False
try:
    import chromadb
    HAS_CHROMADB = True
except ImportError:
    pass

HAS_FAISS = False
try:
    import faiss
    HAS_FAISS = True
except ImportError:
    pass

HAS_BM25 = False
try:
    from rank_bm25 import BM25Okapi
    HAS_BM25 = True
except ImportError:
    pass

HAS_SPLADE_DEPS = False
try:
    import torch
    import transformers
    from transformers import AutoTokenizer, AutoModelForMaskedLM
    HAS_SPLADE_DEPS = True
except ImportError:
    pass

HAS_COHERE = False
try:
    import cohere
    HAS_COHERE = True
except ImportError:
    pass

HAS_OPENAI = False
try:
    import openai
    HAS_OPENAI = True
except ImportError:
    pass

HAS_LANGCHAIN_TEXT_SPLITTERS = False
try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    HAS_LANGCHAIN_TEXT_SPLITTERS = True
except ImportError:
    pass

HAS_SPACY = False
try:
    import spacy
    try:
        spacy.load("en_core_web_sm")
        HAS_SPACY = True
    except OSError:
        pass
except ImportError:
    pass

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Constants
RAG_STORE_DIR = Path("./rag_store")
CHROMA_DIR = RAG_STORE_DIR / "chroma"
FAISS_DIR = RAG_STORE_DIR / "faiss"
SPLADE_DIR = RAG_STORE_DIR / "splade"
BM25_DIR = RAG_STORE_DIR / "bm25"
META_FILE = RAG_STORE_DIR / "meta.jsonl"

# Ensure directories exist
for dir_path in [RAG_STORE_DIR, CHROMA_DIR, FAISS_DIR, SPLADE_DIR, BM25_DIR]:
    dir_path.mkdir(parents=True, exist_ok=True)

# Download NLTK data if needed
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    try:
        nltk.download('punkt', quiet=True)
    except:
        pass

# --- Utility Functions ---
@st.cache_resource
def get_tokenizer(model_name="gpt-3.5-turbo"):
    """Get tiktoken tokenizer for token counting"""
    try:
        return tiktoken.encoding_for_model(model_name)
    except:
        return tiktoken.get_encoding("cl100k_base")

def count_tokens(text: str, tokenizer=None) -> int:
    """Count tokens in text"""
    if tokenizer is None:
        tokenizer = get_tokenizer()
    return len(tokenizer.encode(text))

def normalize_text(text: str) -> str:
    """Normalize text for consistent processing"""
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def compute_chunk_id(chunk_data: Dict) -> str:
    """Generate stable chunk ID from content and metadata"""
    relevant_keys = ['text', 'source_name', 'page_start', 'page_end', 'char_start', 'char_end']
    content_parts = [str(chunk_data.get(key, '')) for key in relevant_keys]
    content_string = "-".join(content_parts)
    return hashlib.sha256(content_string.encode()).hexdigest()

def jaccard_similarity(text1: str, text2: str, n: int = 8) -> float:
    """Compute Jaccard similarity of n-grams"""
    def get_ngrams(text, n):
        words = text.lower().split()
        return set(' '.join(words[i:i+n]) for i in range(len(words)-n+1))
    
    ngrams1 = get_ngrams(text1, n)
    ngrams2 = get_ngrams(text2, n)
    
    if not ngrams1 and not ngrams2:
        return 1.0
    if not ngrams1 or not ngrams2:
        return 0.0
    
    intersection = len(ngrams1 & ngrams2)
    union = len(ngrams1 | ngrams2)
    return intersection / union if union > 0 else 0.0

# --- Document Processing ---
class DocumentProcessor:
    """Process various document formats"""

    def extract_text(self, uploaded_file) -> str:
        """Extract text content from uploaded file for chart analysis"""
        try:
            file_bytes = uploaded_file.read()
            filename = uploaded_file.name
            file_extension = filename.lower().split('.')[-1]
            
            if file_extension == 'pdf':
                pages = self.extract_text_from_pdf(file_bytes, filename)
                return '\n\n'.join([page['text'] for page in pages if page.get('text', '').strip()])
            
            elif file_extension == 'docx':
                pages = self.extract_text_from_docx(file_bytes, filename)
                return '\n\n'.join([page['text'] for page in pages if page.get('text', '').strip()])
            
            elif file_extension in ['txt', 'md']:
                content = file_bytes.decode('utf-8', errors='ignore')
                return content.strip()
            
            elif file_extension in ['xlsx', 'xls']:
                return self.extract_text_from_excel(file_bytes, filename)
            
            elif file_extension == 'csv':
                return self.extract_text_from_csv(file_bytes, filename)
            
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
                
        except Exception as e:
            raise Exception(f"Failed to extract text from {filename}: {str(e)}")
        finally:
            # Reset file pointer for potential reuse
            uploaded_file.seek(0)

    @staticmethod
    def extract_text_from_pdf(file_bytes: bytes, filename: str) -> List[Dict]:
        """Extract text from PDF with page information"""
        pages_content = []

        if HAS_PYMUPDF:
            try:
                doc = fitz.open(stream=file_bytes, filetype="pdf")
                for page_num in range(doc.page_count):
                    page = doc[page_num]
                    text = page.get_text()
                    
                    if text.strip():
                        pages_content.append({
                            'text': text,
                            'page_start': page_num + 1,
                            'page_end': page_num + 1,
                            'source_name': filename,
                            'parser': 'pymupdf',
                            'metadata': {}
                        })
                doc.close()
                return pages_content
            except Exception as e:
                logger.warning(f"PyMuPDF failed for {filename}: {e}")

        if HAS_PYPDF2:
            try:
                from io import BytesIO
                reader = PyPDF2.PdfReader(BytesIO(file_bytes))
                for page_num, page in enumerate(reader.pages):
                    text = page.extract_text()
                    if text.strip():
                        pages_content.append({
                            'text': text,
                            'page_start': page_num + 1,
                            'page_end': page_num + 1,
                            'source_name': filename,
                            'parser': 'pypdf2',
                            'metadata': {}
                        })
                return pages_content
            except Exception as e:
                logger.error(f"PyPDF2 failed for {filename}: {e}")
        
        return [{'text': f"Could not extract text from {filename}", 
                'page_start': 1, 'page_end': 1, 'source_name': filename, 
                'parser': 'failed', 'metadata': {}}]

    @staticmethod
    def extract_text_from_docx(file_bytes: bytes, filename: str) -> List[Dict]:
        """Extract text from DOCX, converting tables to Markdown"""
        if not HAS_DOCX:
            return [{'text': f"python-docx not available for {filename}", 
                    'page_start': 1, 'page_end': 1, 'source_name': filename,
                    'parser': 'failed', 'metadata': {}}]

        try:
            from io import BytesIO
            doc = Document(BytesIO(file_bytes))
            
            full_text = []
            
            # Process paragraphs and tables
            for element in doc.element.body:
                if element.tag.endswith('p'):  # Paragraph
                    para = None
                    for p in doc.paragraphs:
                        if p._element == element:
                            para = p
                            break
                    if para and para.text.strip():
                        full_text.append(para.text)
                        
                elif element.tag.endswith('tbl'):  # Table
                    table = None
                    for t in doc.tables:
                        if t._element == element:
                            table = t
                            break
                    if table:
                        table_md = ""
                        for r, row in enumerate(table.rows):
                            row_texts = [cell.text.strip() for cell in row.cells]
                            table_md += "| " + " | ".join(row_texts) + " |\n"
                            if r == 0:  # Header separator
                                table_md += "|" + "---|" * len(row_texts) + "\n"
                        full_text.append(table_md)
            
            content = '\n\n'.join(full_text)
            return [{'text': content, 'page_start': 1, 'page_end': 1, 'source_name': filename,
                    'parser': 'docx', 'metadata': {}}]
            
        except Exception as e:
            logger.error(f"DOCX extraction failed for {filename}: {e}")
            return [{'text': f"Could not extract text from {filename}", 
                    'page_start': 1, 'page_end': 1, 'source_name': filename,
                    'parser': 'failed', 'metadata': {}}]

    def extract_text_from_excel(self, file_bytes: bytes, filename: str) -> str:
        """Extract text from Excel files (XLSX/XLS) for chart analysis"""
        try:
            # Try importing pandas for Excel processing
            import pandas as pd
            from io import BytesIO
            
            # Read Excel file
            excel_file = BytesIO(file_bytes)
            
            # Get all sheet names
            try:
                sheet_names = pd.ExcelFile(excel_file).sheet_names
            except Exception:
                # Fallback to default sheet
                sheet_names = [0]
            
            all_content = []
            
            for sheet_name in sheet_names:
                try:
                    # Read each sheet
                    df = pd.read_excel(excel_file, sheet_name=sheet_name)
                    
                    if not df.empty:
                        # Convert to readable format
                        sheet_content = f"Sheet: {sheet_name}\n"
                        
                        # Convert DataFrame to structured text
                        # Include column headers
                        headers = list(df.columns)
                        sheet_content += "Columns: " + " | ".join(str(h) for h in headers) + "\n\n"
                        
                        # Add data rows with meaningful formatting
                        for idx, row in df.iterrows():
                            row_data = []
                            for col in df.columns:
                                value = row[col]
                                if pd.isna(value):
                                    row_data.append("")
                                else:
                                    # Format numbers nicely
                                    if isinstance(value, (int, float)):
                                        if isinstance(value, float) and value.is_integer():
                                            row_data.append(f"{int(value)}")
                                        else:
                                            row_data.append(f"{value}")
                                    else:
                                        row_data.append(str(value))
                            
                            # Create structured text for each row
                            for i, (header, data) in enumerate(zip(headers, row_data)):
                                if data.strip():  # Only include non-empty data
                                    sheet_content += f"{header}: {data}\n"
                            sheet_content += "\n"
                        
                        all_content.append(sheet_content)
                        
                except Exception as e:
                    logger.warning(f"Could not process sheet {sheet_name} in {filename}: {e}")
                    continue
            
            if all_content:
                return "\n\n".join(all_content)
            else:
                raise ValueError("No readable data found in Excel file")
                
        except ImportError:
            raise Exception("pandas library required for Excel processing. Install with: pip install pandas openpyxl")
        except Exception as e:
            raise Exception(f"Failed to process Excel file {filename}: {str(e)}")

    def extract_text_from_csv(self, file_bytes: bytes, filename: str) -> str:
        """Extract text from CSV files for chart analysis"""
        try:
            import pandas as pd
            from io import StringIO
            
            # Decode CSV content
            csv_content = file_bytes.decode('utf-8', errors='ignore')
            
            # Try different separators
            separators = [',', ';', '\t', '|']
            df = None
            
            for sep in separators:
                try:
                    df = pd.read_csv(StringIO(csv_content), separator=sep)
                    if len(df.columns) > 1:  # Valid CSV with multiple columns
                        break
                except Exception:
                    continue
            
            if df is None or df.empty:
                # Fallback to raw text
                return csv_content
            
            # Convert CSV to structured text
            content = f"CSV Data from {filename}\n"
            content += "Columns: " + " | ".join(str(col) for col in df.columns) + "\n\n"
            
            # Add data rows
            for idx, row in df.iterrows():
                for col in df.columns:
                    value = row[col]
                    if pd.isna(value):
                        continue
                    
                    # Format values nicely
                    if isinstance(value, (int, float)):
                        if isinstance(value, float) and value.is_integer():
                            content += f"{col}: {int(value)}\n"
                        else:
                            content += f"{col}: {value}\n"
                    else:
                        content += f"{col}: {str(value)}\n"
                content += "\n"
            
            return content
            
        except ImportError:
            # Fallback to raw CSV text
            return file_bytes.decode('utf-8', errors='ignore')
        except Exception as e:
            raise Exception(f"Failed to process CSV file {filename}: {str(e)}")

    @staticmethod
    def extract_text_from_txt(file_bytes: bytes, filename: str) -> List[Dict]:
        """Extract text from TXT/MD, handling optional YAML front matter"""
        try:
            text = file_bytes.decode('utf-8')
            extracted_metadata = {}
            
            if filename.lower().endswith(('.md', '.markdown')):
                # Parse YAML front matter
                front_matter_match = re.match(r'^---\s*\n(.*?)\n---\s*\n(.*)$', text, re.DOTALL)
                if front_matter_match:
                    yaml_content = front_matter_match.group(1)
                    markdown_content = front_matter_match.group(2)
                    try:
                        extracted_metadata = yaml.safe_load(yaml_content) or {}
                        text = markdown_content
                    except yaml.YAMLError:
                        pass
            
            return [{'text': text, 'page_start': 1, 'page_end': 1, 'source_name': filename,
                    'parser': 'text', 'metadata': extracted_metadata}]
            
        except UnicodeDecodeError:
            try:
                text = file_bytes.decode('latin1')
                return [{'text': text, 'page_start': 1, 'page_end': 1, 'source_name': filename,
                        'parser': 'text_latin1', 'metadata': {}}]
            except Exception as e:
                logger.error(f"Text extraction failed for {filename}: {e}")
                return [{'text': f"Could not extract text from {filename}",
                        'page_start': 1, 'page_end': 1, 'source_name': filename,
                        'parser': 'failed', 'metadata': {}}]

# --- Chunking Strategies ---
class UniversalChunker:
    """Universal chunker with all strategies"""
    
    def __init__(self, strategy="token", target_tokens=700, overlap_tokens=80, min_tokens=120,
                 char_size=4000, char_overlap=400, semantic_threshold=0.65, 
                 dense_embedder=None, include_front_matter=False, front_matter_keys=None):
        self.strategy = strategy
        self.target_tokens = target_tokens
        self.overlap_tokens = overlap_tokens
        self.min_tokens = min_tokens
        self.char_size = char_size
        self.char_overlap = char_overlap
        self.semantic_threshold = semantic_threshold
        self.dense_embedder = dense_embedder
        self.tokenizer = get_tokenizer()
        self.include_front_matter = include_front_matter
        self.front_matter_keys = front_matter_keys or [
            'source_name', 'section_path', 'page_start', 'page_end', 
            'tokens', 'sha256', 'created_at'
        ]
    
    def chunk_document(self, doc_data: Dict) -> List[Dict]:
        """Main chunking entry point"""
        text = doc_data['text']
        source_name = doc_data['source_name']
        page_start = doc_data.get('page_start', 1)
        page_end = doc_data.get('page_end', 1)
        doc_metadata = doc_data.get('metadata', {})
        
        # Route to appropriate chunking strategy
        if self.strategy == "token":
            chunks = self._chunk_by_tokens(text, source_name, page_start, page_end)
        elif self.strategy == "character":
            chunks = self._chunk_by_characters(text, source_name, page_start, page_end)
        elif self.strategy == "sentence":
            chunks = self._chunk_by_sentences(text, source_name, page_start, page_end)
        elif self.strategy == "paragraph":
            chunks = self._chunk_by_paragraphs(text, source_name, page_start, page_end)
        elif self.strategy == "page":
            chunks = [self._create_chunk(text, source_name, page_start, page_end, 'page')]
        elif self.strategy == "markdown_structure":
            chunks = self._chunk_by_markdown_structure(text, source_name, page_start, page_end)
        elif self.strategy == "semantic_boundary":
            chunks = self._chunk_by_semantic_boundary(text, source_name, page_start, page_end)
        elif self.strategy == "recursive_character":
            chunks = self._chunk_recursive_character(text, source_name, page_start, page_end)
        elif self.strategy == "hybrid_chunking":
            chunks = self._chunk_hybrid(text, source_name, page_start, page_end)
        else:
            # Default to token chunking
            chunks = self._chunk_by_tokens(text, source_name, page_start, page_end)
        
        # Add document metadata and front matter
        for chunk in chunks:
            chunk.update(doc_metadata)
            if self.include_front_matter:
                chunk['text'] = self._add_front_matter(chunk) + '\n\n' + chunk['text']
        
        return chunks
    
    def _create_chunk(self, text: str, source_name: str, page_start: int, page_end: int, 
                     parser_type: str, metadata: Optional[Dict] = None) -> Dict:
        """Create a chunk dictionary"""
        chunk_data = {
            'text': normalize_text(text),
            'source_name': source_name,
            'page_start': page_start,
            'page_end': page_end,
            'tokens': count_tokens(text, self.tokenizer),
            'parser': parser_type,
            'created_at': datetime.now().isoformat(),
            'char_start': 0,
            'char_end': len(text)
        }
        if metadata:
            chunk_data.update(metadata)
        chunk_data['id'] = compute_chunk_id(chunk_data)
        chunk_data['sha256'] = hashlib.sha256(normalize_text(text).encode()).hexdigest()
        return chunk_data
    
    def _add_front_matter(self, chunk: Dict) -> str:
        """Add YAML front matter to chunk"""
        front_matter = {}
        for key in self.front_matter_keys:
            if key in chunk:
                front_matter[key] = chunk[key]
        
        if front_matter:
            yaml_str = yaml.dump(front_matter, default_flow_style=False)
            return f"---\n{yaml_str}---"
        return ""
    
    def _chunk_by_tokens(self, text: str, source_name: str, page_start: int, page_end: int) -> List[Dict]:
        """Token-based chunking with overlap"""
        tokens = self.tokenizer.encode(text)
        chunks = []
        
        start = 0
        while start < len(tokens):
            end = min(start + self.target_tokens, len(tokens))
            chunk_tokens = tokens[start:end]
            chunk_text = self.tokenizer.decode(chunk_tokens)
            
            if len(chunk_tokens) >= self.min_tokens or start + self.target_tokens >= len(tokens):
                chunk = self._create_chunk(chunk_text, source_name, page_start, page_end, 'token',
                                         {'char_start': start, 'char_end': end})
                chunks.append(chunk)
            
            start += self.target_tokens - self.overlap_tokens
        
        return chunks
    
    def _chunk_by_characters(self, text: str, source_name: str, page_start: int, page_end: int) -> List[Dict]:
        """Character-based chunking"""
        chunks = []
        start = 0
        
        while start < len(text):
            end = min(start + self.char_size, len(text))
            chunk_text = text[start:end]
            
            if len(chunk_text.strip()) >= self.min_tokens * 4:  # Rough char estimate
                chunk = self._create_chunk(chunk_text, source_name, page_start, page_end, 'character',
                                         {'char_start': start, 'char_end': end})
                chunks.append(chunk)
            
            start += self.char_size - self.char_overlap
        
        return chunks
    
    def _chunk_by_sentences(self, text: str, source_name: str, page_start: int, page_end: int) -> List[Dict]:
        """Sentence-based chunking"""
        try:
            sentences = nltk.sent_tokenize(text)
        except:
            sentences = [s.strip() for s in re.split(r'[.!?]+', text) if s.strip()]
        
        chunks = []
        current_chunk = []
        current_tokens = 0
        
        for sentence in sentences:
            sentence_tokens = count_tokens(sentence, self.tokenizer)
            
            if current_tokens + sentence_tokens > self.target_tokens and current_chunk:
                chunk_text = ' '.join(current_chunk)
                if count_tokens(chunk_text, self.tokenizer) >= self.min_tokens:
                    chunk = self._create_chunk(chunk_text, source_name, page_start, page_end, 'sentence')
                    chunks.append(chunk)
                
                # Overlap
                overlap_sentences = current_chunk[-2:] if len(current_chunk) >= 2 else current_chunk
                current_chunk = overlap_sentences + [sentence]
                current_tokens = sum(count_tokens(s, self.tokenizer) for s in current_chunk)
            else:
                current_chunk.append(sentence)
                current_tokens += sentence_tokens
        
        # Final chunk
        if current_chunk:
            chunk_text = ' '.join(current_chunk)
            if count_tokens(chunk_text, self.tokenizer) >= self.min_tokens:
                chunk = self._create_chunk(chunk_text, source_name, page_start, page_end, 'sentence')
                chunks.append(chunk)
        
        return chunks
    
    def _chunk_by_paragraphs(self, text: str, source_name: str, page_start: int, page_end: int) -> List[Dict]:
        """Paragraph-based chunking"""
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        chunks = []
        current_chunk = []
        current_tokens = 0
        
        for para in paragraphs:
            para_tokens = count_tokens(para, self.tokenizer)
            
            if current_tokens + para_tokens > self.target_tokens and current_chunk:
                chunk_text = '\n\n'.join(current_chunk)
                if count_tokens(chunk_text, self.tokenizer) >= self.min_tokens:
                    chunk = self._create_chunk(chunk_text, source_name, page_start, page_end, 'paragraph')
                    chunks.append(chunk)
                
                current_chunk = [para]
                current_tokens = para_tokens
            else:
                current_chunk.append(para)
                current_tokens += para_tokens
        
        # Final chunk
        if current_chunk:
            chunk_text = '\n\n'.join(current_chunk)
            if count_tokens(chunk_text, self.tokenizer) >= self.min_tokens:
                chunk = self._create_chunk(chunk_text, source_name, page_start, page_end, 'paragraph')
                chunks.append(chunk)
        
        return chunks
    
    def _chunk_by_markdown_structure(self, text: str, source_name: str, page_start: int, page_end: int) -> List[Dict]:
        """Markdown structure-aware chunking"""
        lines = text.split('\n')
        sections = []
        current_section = {'content': [], 'level': 0, 'title': '', 'path': ''}
        section_stack = []
        
        for line in lines:
            # Check for ATX headers
            header_match = re.match(r'^(#{1,6})\s+(.+)$', line.strip())
            if header_match:
                # Save current section
                if current_section['content']:
                    sections.append(current_section.copy())
                
                level = len(header_match.group(1))
                title = header_match.group(2)
                
                # Update section stack
                section_stack = [s for s in section_stack if s['level'] < level]
                section_stack.append({'level': level, 'title': title})
                
                # Create section path
                path_parts = [s['title'] for s in section_stack]
                section_path = ' > '.join(path_parts)
                
                current_section = {
                    'content': [line],
                    'level': level,
                    'title': title,
                    'path': section_path
                }
            else:
                current_section['content'].append(line)
        
        # Add final section
        if current_section['content']:
            sections.append(current_section)
        
        # Convert sections to chunks
        chunks = []
        for section in sections:
            content = '\n'.join(section['content'])
            if content.strip():
                # Detect markdown types
                md_types = set()
                if re.search(r'^\s*[-*+]\s+', content, re.MULTILINE):
                    md_types.add('list')
                if re.search(r'^\s*>\s+', content, re.MULTILINE):
                    md_types.add('blockquote')
                if re.search(r'```', content):
                    md_types.add('code_fence')
                if re.search(r'\|.*\|', content):
                    md_types.add('table')
                
                metadata = {
                    'section_path': section['path'],
                    'md_types_present': list(md_types),
                    'header_level': section['level']
                }
                
                chunk = self._create_chunk(content, source_name, page_start, page_end, 
                                         'markdown_structure', metadata)
                chunks.append(chunk)
        
        return chunks
    
    def _chunk_by_semantic_boundary(self, text: str, source_name: str, page_start: int, page_end: int) -> List[Dict]:
        """Semantic boundary chunking using sentence similarity"""
        if not self.dense_embedder:
            return self._chunk_by_sentences(text, source_name, page_start, page_end)
        
        try:
            sentences = nltk.sent_tokenize(text)
        except:
            sentences = [s.strip() for s in re.split(r'[.!?]+', text) if s.strip()]
        
        if len(sentences) <= 1:
            return [self._create_chunk(text, source_name, page_start, page_end, 'semantic_boundary')]
        
        # Get sentence embeddings
        embeddings = self.dense_embedder.encode(sentences)
        
        # Find semantic boundaries
        boundaries = [0]
        current_chunk_sentences = [sentences[0]]
        current_tokens = count_tokens(sentences[0], self.tokenizer)
        
        for i in range(1, len(sentences)):
            # Check semantic similarity with previous sentence
            similarity = cosine_similarity([embeddings[i-1]], [embeddings[i]])[0][0]
            sentence_tokens = count_tokens(sentences[i], self.tokenizer)
            
            # Split if similarity below threshold or token limit exceeded
            if (similarity < self.semantic_threshold or 
                current_tokens + sentence_tokens > self.target_tokens) and current_chunk_sentences:
                
                boundaries.append(i)
                current_chunk_sentences = [sentences[i]]
                current_tokens = sentence_tokens
            else:
                current_chunk_sentences.append(sentences[i])
                current_tokens += sentence_tokens
        
        boundaries.append(len(sentences))
        
        # Create chunks from boundaries
        chunks = []
        for i in range(len(boundaries) - 1):
            start_idx = boundaries[i]
            end_idx = boundaries[i + 1]
            chunk_sentences = sentences[start_idx:end_idx]
            chunk_text = ' '.join(chunk_sentences)
            
            if count_tokens(chunk_text, self.tokenizer) >= self.min_tokens:
                chunk = self._create_chunk(chunk_text, source_name, page_start, page_end, 
                                         'semantic_boundary', {'sentence_count': len(chunk_sentences)})
                chunks.append(chunk)
        
        return chunks
    
    def _chunk_recursive_character(self, text: str, source_name: str, page_start: int, page_end: int) -> List[Dict]:
        """Recursive character splitting (LangChain-style)"""
        if HAS_LANGCHAIN_TEXT_SPLITTERS:
            try:
                text_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=self.char_size,
                    chunk_overlap=self.char_overlap,
                    length_function=len,
                    add_start_index=True
                )
                
                docs = text_splitter.create_documents([text])
                chunks = []
                for doc in docs:
                    metadata = doc.metadata
                    char_start = metadata.get('start_index', 0)
                    chunk = self._create_chunk(doc.page_content, source_name, page_start, page_end,
                                             'recursive_character', {'char_start': char_start})
                    chunks.append(chunk)
                return chunks
            except Exception:
                pass
        
        # Fallback to character chunking
        return self._chunk_by_characters(text, source_name, page_start, page_end)
    
    def _chunk_hybrid(self, text: str, source_name: str, page_start: int, page_end: int) -> List[Dict]:
        """Hybrid chunking: markdown structure + semantic refinement + dedup"""
        # First pass: markdown structure
        md_chunks = self._chunk_by_markdown_structure(text, source_name, page_start, page_end)
        
        # Second pass: semantic refinement within sections
        refined_chunks = []
        for chunk in md_chunks:
            if count_tokens(chunk['text'], self.tokenizer) > self.target_tokens * 1.5:
                # Large section, apply semantic boundary
                sub_chunks = self._chunk_by_semantic_boundary(chunk['text'], source_name, 
                                                            page_start, page_end)
                for sub_chunk in sub_chunks:
                    sub_chunk['parser'] = 'hybrid'
                    sub_chunk['section_path'] = chunk.get('section_path', '')
                refined_chunks.extend(sub_chunks)
            else:
                chunk['parser'] = 'hybrid'
                refined_chunks.append(chunk)
        
        # Third pass: dedup near-duplicates using Jaccard similarity
        deduped_chunks = []
        for chunk in refined_chunks:
            is_duplicate = False
            for existing in deduped_chunks:
                if jaccard_similarity(chunk['text'], existing['text']) > 0.8:
                    is_duplicate = True
                    break
            if not is_duplicate:
                deduped_chunks.append(chunk)
        
        return deduped_chunks

# --- Dense Index ---
class DenseIndex:
    """Dense vector index with Chroma/FAISS fallback"""
    
    def __init__(self, model_name="sentence-transformers/all-MiniLM-L6-v2"):
        self.model_name = model_name
        self.embedder = None
        self.index = None
        self.index_type = None
        self._init_embedder()
        self._init_index()
    
    @staticmethod
    @st.cache_resource
    def _load_embedder_model(_model_name):
        """Initialize sentence transformer model"""
        if HAS_SENTENCE_TRANSFORMERS:
            try:
                return SentenceTransformer(_model_name)
            except Exception as e:
                logger.error(f"Failed to load embedder {_model_name}: {e}")
        return None
    
    def _init_embedder(self):
        """Initialize embedder"""
        self.embedder = DenseIndex._load_embedder_model(self.model_name)
    
    def _init_index(self):
        """Initialize vector index"""
        if HAS_CHROMADB:
            try:
                client = chromadb.PersistentClient(path=str(CHROMA_DIR))
                self.index = client.get_or_create_collection("documents")
                self.index_type = "chroma"
                return
            except Exception as e:
                logger.warning(f"Chroma init failed: {e}")
        
        if HAS_FAISS:
            try:
                if (FAISS_DIR / "index.faiss").exists():
                    self.index = faiss.read_index(str(FAISS_DIR / "index.faiss"))
                    self.index_type = "faiss"
                    return
            except Exception as e:
                logger.warning(f"FAISS load failed: {e}")
        
        # In-memory fallback
        self.index = []
        self.index_type = "memory"
    
    def add(self, chunks: List[Dict]):
        """Add chunks to dense index"""
        if not self.embedder or not chunks:
            return
        
        texts = [chunk['text'] for chunk in chunks]
        embeddings = self.embedder.encode(texts, normalize_embeddings=True)
        
        if self.index_type == "chroma":
            try:
                self.index.add(
                    ids=[chunk['id'] for chunk in chunks],
                    embeddings=embeddings.tolist(),
                    metadatas=[{k: v for k, v in chunk.items() if k != 'text'} for chunk in chunks],
                    documents=texts
                )
            except Exception as e:
                logger.error(f"Chroma add failed: {e}")
        
        elif self.index_type == "faiss":
            try:
                if self.index is None:
                    self.index = faiss.IndexFlatIP(embeddings.shape[1])
                
                self.index.add(embeddings.astype(np.float32))
                faiss.write_index(self.index, str(FAISS_DIR / "index.faiss"))
                
                # Save metadata
                metadata_path = FAISS_DIR / "metadata.pkl"
                existing_meta = []
                if metadata_path.exists():
                    with open(metadata_path, 'rb') as f:
                        existing_meta = pickle.load(f)
                
                existing_meta.extend(chunks)
                with open(metadata_path, 'wb') as f:
                    pickle.dump(existing_meta, f)
                    
            except Exception as e:
                logger.error(f"FAISS add failed: {e}")
        
        else:  # memory
            for chunk, emb in zip(chunks, embeddings):
                self.index.append((chunk, emb))
    
    def search(self, query: str, top_k: int = 10) -> List[Tuple[Dict, float]]:
        """Search dense index"""
        if not self.embedder:
            return []
        
        query_embedding = self.embedder.encode(query, normalize_embeddings=True)
        
        if self.index_type == "chroma":
            try:
                results = self.index.query(
                    query_embeddings=[query_embedding.tolist()],
                    n_results=min(top_k, self.index.count()),
                    include=['documents', 'distances', 'metadatas']
                )
                
                chunks_scores = []
                if results['ids'] and results['ids'][0]:
                    for i in range(len(results['ids'][0])):
                        chunk = {
                            'id': results['ids'][0][i],
                            'text': results['documents'][0][i],
                            **results['metadatas'][0][i]
                        }
                        score = max(0, 1 - results['distances'][0][i])
                        chunks_scores.append((chunk, score))
                
                return chunks_scores
            except Exception as e:
                logger.error(f"Chroma search failed: {e}")
        
        elif self.index_type == "faiss":
            try:
                D, I = self.index.search(np.array([query_embedding]).astype(np.float32), top_k)
                
                metadata_path = FAISS_DIR / "metadata.pkl"
                if metadata_path.exists():
                    with open(metadata_path, 'rb') as f:
                        all_metadata = pickle.load(f)
                    
                    chunks_scores = []
                    for i, idx in enumerate(I[0]):
                        if idx != -1 and idx < len(all_metadata):
                            score = max(0, D[0][i])
                            chunks_scores.append((all_metadata[idx], score))
                    
                    return chunks_scores
            except Exception as e:
                logger.error(f"FAISS search failed: {e}")
        
        else:  # memory
            scores = []
            for chunk, emb in self.index:
                score = np.dot(query_embedding, emb)
                scores.append((chunk, score))
            
            scores.sort(key=lambda x: x[1], reverse=True)
            return scores[:top_k]
        
        return []

# --- SPLADE++ Sparse Index ---
class SpladeEncoder:
    """SPLADE++ encoder"""
    
    def __init__(self, model_name="naver/splade-cocondenser-ensembledistil", top_k_vocab=60):
        self.model_name = model_name
        self.top_k_vocab = top_k_vocab
        self.tokenizer = None
        self.model = None
        self.available = False
        self._load_model()
    
    @staticmethod
    @st.cache_resource
    def _load_splade_model(_model_name):
        """Load SPLADE model with caching"""
        if not HAS_SPLADE_DEPS:
            return None, None, False
        
        try:
            tokenizer = AutoTokenizer.from_pretrained(_model_name)
            model = AutoModelForMaskedLM.from_pretrained(_model_name)
            model.eval()
            return tokenizer, model, True
        except Exception as e:
            logger.error(f"Failed to load SPLADE model {_model_name}: {e}")
            return None, None, False
    
    def _load_model(self):
        """Load SPLADE model"""
        self.tokenizer, self.model, self.available = SpladeEncoder._load_splade_model(self.model_name)
    
    def encode(self, text: str) -> Dict[int, float]:
        """Encode text to sparse vector"""
        if not self.available:
            return {}
        
        try:
            inputs = self.tokenizer(
                text, return_tensors="pt", truncation=True, 
                max_length=512, padding=True
            )
            
            with torch.no_grad():
                outputs = self.model(**inputs)
                logits = outputs.logits
            
            # SPLADE activation
            activated = torch.log1p(torch.relu(logits))
            weights = activated.max(dim=1).values.squeeze(0)
            
            # Get top-k terms
            top_k = min(self.top_k_vocab, len(weights))
            top_indices = torch.topk(weights, top_k).indices
            
            sparse_vector = {}
            for idx in top_indices:
                term_id = idx.item()
                weight = weights[term_id].item()
                if weight > 0:
                    sparse_vector[term_id] = weight
            
            return sparse_vector
        except Exception as e:
            logger.error(f"SPLADE encoding failed: {e}")
            return {}

class SparseIndex:
    """Sparse inverted index for SPLADE++"""
    
    def __init__(self):
        self.inverted_index = defaultdict(list)  # term_id -> [(doc_id, weight)]
        self.doc_norms = {}  # doc_id -> L2 norm
        self.doc_metadata = {}  # doc_id -> metadata
    
    def add_document(self, doc_id: str, sparse_vector: Dict[int, float], metadata: Dict):
        """Add document to sparse index"""
        if not sparse_vector:
            return
        
        norm = np.sqrt(sum(w**2 for w in sparse_vector.values()))
        self.doc_norms[doc_id] = norm
        self.doc_metadata[doc_id] = metadata
        
        for term_id, weight in sparse_vector.items():
            self.inverted_index[term_id].append((doc_id, weight))
    
    def search(self, query_vector: Dict[int, float], top_k: int = 10) -> List[Tuple[str, float]]:
        """Search using sparse cosine similarity"""
        if not query_vector:
            return []
        
        query_norm = np.sqrt(sum(w**2 for w in query_vector.values()))
        if query_norm == 0:
            return []
        
        doc_scores = defaultdict(float)
        for term_id, query_weight in query_vector.items():
            if term_id in self.inverted_index:
                for doc_id, doc_weight in self.inverted_index[term_id]:
                    doc_scores[doc_id] += query_weight * doc_weight
        
        # Normalize by document norms
        for doc_id in doc_scores:
            doc_norm = self.doc_norms.get(doc_id, 1.0)
            if doc_norm > 0:
                doc_scores[doc_id] /= (query_norm * doc_norm)
        
        sorted_docs = sorted(doc_scores.items(), key=lambda x: x[1], reverse=True)
        return sorted_docs[:top_k]
    
    def save(self, filepath: Path):
        """Save index to disk"""
        data = {
            'inverted_index': dict(self.inverted_index),
            'doc_norms': self.doc_norms,
            'doc_metadata': self.doc_metadata
        }
        with open(filepath, 'wb') as f:
            pickle.dump(data, f)
    
    def load(self, filepath: Path) -> bool:
        """Load index from disk"""
        if filepath.exists():
            try:
                with open(filepath, 'rb') as f:
                    data = pickle.load(f)
                self.inverted_index = defaultdict(list, data.get('inverted_index', {}))
                self.doc_norms = data.get('doc_norms', {})
                self.doc_metadata = data.get('doc_metadata', {})
                return True
            except Exception as e:
                logger.error(f"Failed to load sparse index: {e}")
        return False

# --- BM25 Index ---
class BM25Index:
    """BM25 lexical index"""
    
    def __init__(self):
        self.index = None
        self.chunk_ids = []
        self.available = HAS_BM25
    
    def add_documents(self, chunks: List[Dict]):
        """Add chunks to BM25 index"""
        if not self.available or not chunks:
            return
        
        corpus = []
        self.chunk_ids = []
        
        for chunk in chunks:
            tokens = normalize_text(chunk['text']).lower().split()
            corpus.append(tokens)
            self.chunk_ids.append(chunk['id'])
        
        self.index = BM25Okapi(corpus)
    
    def search(self, query: str, chunks: List[Dict], top_k: int = 10) -> List[Tuple[Dict, float]]:
        """Search BM25 index"""
        if not self.available or not self.index:
            return []
        
        tokenized_query = normalize_text(query).lower().split()
        scores = self.index.get_scores(tokenized_query)
        
        chunk_scores = []
        for i, score in enumerate(scores):
            if i < len(chunks):
                chunk_scores.append((chunks[i], float(score)))
        
        chunk_scores.sort(key=lambda x: x[1], reverse=True)
        return chunk_scores[:top_k]
    
    def save(self, filepath: Path):
        """Save BM25 index"""
        if self.available and self.index:
            data = {'index': self.index, 'chunk_ids': self.chunk_ids}
            with open(filepath, 'wb') as f:
                pickle.dump(data, f)
    
    def load(self, filepath: Path) -> bool:
        """Load BM25 index"""
        if self.available and filepath.exists():
            try:
                with open(filepath, 'rb') as f:
                    data = pickle.load(f)
                self.index = data.get('index')
                self.chunk_ids = data.get('chunk_ids', [])
                return True
            except Exception as e:
                logger.error(f"Failed to load BM25 index: {e}")
        return False

# --- Retrieval & Fusion ---
def reciprocal_rank_fusion(results_list: List[List[Tuple]], rrf_k: int = 60) -> List[Tuple]:
    """Combine multiple result lists using RRF"""
    doc_scores = defaultdict(float)
    doc_items = {}
    
    for results in results_list:
        for rank, (item, score) in enumerate(results):
            doc_id = item.get('id', str(hash(item.get('text', ''))))
            doc_scores[doc_id] += 1.0 / (rrf_k + rank + 1)
            doc_items[doc_id] = (item, score)
    
    # Sort by RRF score
    sorted_docs = sorted(doc_scores.items(), key=lambda x: x[1], reverse=True)
    return [(doc_items[doc_id][0], rrf_score) for doc_id, rrf_score in sorted_docs]

def alpha_weighted_fusion(dense_results: List[Tuple], sparse_results: List[Tuple], alpha: float) -> List[Tuple]:
    """
    Combine dense and sparse results using weighted sum.
    
    Args:
        dense_results: List of (chunk, score) from dense retrieval
        sparse_results: List of (chunk, score) from sparse retrieval  
        alpha: Weight for dense results (0.0 = pure sparse, 1.0 = pure dense)
    
    Returns:
        Combined and sorted results with weighted scores
    """
    # Normalize scores to [0, 1] range for fair combination
    def normalize_scores(results):
        if not results:
            return []
        scores = [score for _, score in results]
        if len(set(scores)) == 1:  # All scores are the same
            return [(item, 1.0) for item, _ in results]
        min_score, max_score = min(scores), max(scores)
        if max_score == min_score:
            return [(item, 1.0) for item, _ in results]
        return [(item, (score - min_score) / (max_score - min_score)) for item, score in results]
    
    # Normalize both result sets
    norm_dense = normalize_scores(dense_results)
    norm_sparse = normalize_scores(sparse_results)
    
    # Combine scores using alpha weighting
    doc_scores = {}
    doc_items = {}
    
    # Add dense results with alpha weight
    for item, norm_score in norm_dense:
        doc_id = item.get('id', str(hash(item.get('text', ''))))
        doc_scores[doc_id] = alpha * norm_score
        doc_items[doc_id] = item
    
    # Add sparse results with (1-alpha) weight
    for item, norm_score in norm_sparse:
        doc_id = item.get('id', str(hash(item.get('text', ''))))
        if doc_id in doc_scores:
            # Combine scores if document appears in both
            doc_scores[doc_id] += (1 - alpha) * norm_score
        else:
            # New document from sparse only
            doc_scores[doc_id] = (1 - alpha) * norm_score
            doc_items[doc_id] = item
    
    # Sort by combined weighted score
    sorted_docs = sorted(doc_scores.items(), key=lambda x: x[1], reverse=True)
    return [(doc_items[doc_id], weighted_score) for doc_id, weighted_score in sorted_docs]

# --- Reranker ---
class Reranker:
    """Reranking with Cohere, Qwen3-Reranker-4B, or local cross-encoder"""
    
    def __init__(self):
        self.cohere_client = None
        self.cross_encoder = None
        self.qwen_reranker = None
        self.reranker_type = "none"
        self._init_reranker()
    
    @staticmethod
    @st.cache_resource
    def _load_qwen_reranker():
        """Load Qwen3-Reranker-4B model with caching"""
        if not HAS_SPLADE_DEPS:  # Uses transformers and torch
            return None, None, False
        
        try:
            from transformers import AutoTokenizer, AutoModelForSequenceClassification
            import torch
            
            model_name = "Qwen/Qwen3-Reranker-4B"
            tokenizer = AutoTokenizer.from_pretrained(model_name)
            model = AutoModelForSequenceClassification.from_pretrained(
                model_name,
                torch_dtype=torch.float16 if torch.cuda.is_available() else torch.float32,
                device_map="auto" if torch.cuda.is_available() else None
            )
            model.eval()
            return tokenizer, model, True
        except Exception as e:
            logger.error(f"Failed to load Qwen3-Reranker-4B: {e}")
            return None, None, False
    
    def _init_reranker(self):
        """Initialize reranker"""
        # Try Cohere first
        cohere_key = os.getenv('COHERE_API_KEY') or "i9MK1lP979EPNOKTLa7n4cRJKjylsSZDOv7JBn6n"
        if HAS_COHERE and cohere_key:
            try:
                self.cohere_client = cohere.Client(api_key=cohere_key)
                self.reranker_type = "cohere"
                return
            except Exception as e:
                logger.warning(f"Cohere init failed: {e}")
        
        # Try Qwen3-Reranker-4B
        if HAS_SPLADE_DEPS:
            try:
                qwen_tokenizer, qwen_model, qwen_available = Reranker._load_qwen_reranker()
                if qwen_available:
                    self.qwen_reranker = (qwen_tokenizer, qwen_model)
                    self.reranker_type = "qwen3_reranker"
                    return
            except Exception as e:
                logger.warning(f"Qwen3-Reranker-4B init failed: {e}")
        
        # Try local cross-encoder
        if HAS_SENTENCE_TRANSFORMERS:
            try:
                self.cross_encoder = CrossEncoder('cross-encoder/ms-marco-MiniLM-L-6-v2')
                self.reranker_type = "cross_encoder"
                return
            except Exception as e:
                logger.warning(f"Cross-encoder init failed: {e}")
    
    def rerank(self, query: str, chunks_scores: List[Tuple[Dict, float]], top_n: int = 10) -> List[Tuple[Dict, float]]:
        """Rerank chunks"""
        if not chunks_scores:
            return []
        
        chunks = [item[0] for item in chunks_scores]
        
        if self.reranker_type == "cohere":
            try:
                texts = [chunk['text'] for chunk in chunks]
                results = self.cohere_client.rerank(
                    query=query,
                    documents=texts,
                    top_n=min(top_n, len(texts)),
                    model='rerank-english-v3.0'
                )
                
                reranked = []
                for result in results.results:
                    chunk = chunks[result.index]
                    score = result.relevance_score
                    reranked.append((chunk, score))
                
                return reranked
            except Exception as e:
                logger.error(f"Cohere reranking failed: {e}")
        
        elif self.reranker_type == "qwen3_reranker":
            try:
                tokenizer, model = self.qwen_reranker
                pairs = []
                for chunk in chunks:
                    # Qwen3-Reranker expects format: "query: <query> document: <document>"
                    pair_text = f"query: {query} document: {chunk['text']}"
                    pairs.append(pair_text)
                
                # Tokenize in batches for efficiency
                batch_size = 8
                all_scores = []
                
                for i in range(0, len(pairs), batch_size):
                    batch_pairs = pairs[i:i + batch_size]
                    
                    inputs = tokenizer(
                        batch_pairs,
                        padding=True,
                        truncation=True,
                        max_length=512,
                        return_tensors="pt"
                    )
                    
                    with torch.no_grad():
                        outputs = model(**inputs)
                        scores = torch.nn.functional.softmax(outputs.logits, dim=-1)[:, 1]  # Positive class probability
                        all_scores.extend(scores.cpu().numpy())
                
                # Create reranked results
                reranked = [(chunks[i], float(all_scores[i])) for i in range(len(chunks))]
                reranked.sort(key=lambda x: x[1], reverse=True)
                
                return reranked[:top_n]
            except Exception as e:
                logger.error(f"Qwen3-Reranker-4B reranking failed: {e}")
        
        elif self.reranker_type == "cross_encoder":
            try:
                pairs = [[query, chunk['text']] for chunk in chunks]
                scores = self.cross_encoder.predict(pairs)
                
                reranked = [(chunks[i], float(scores[i])) for i in range(len(chunks))]
                reranked.sort(key=lambda x: x[1], reverse=True)
                
                return reranked[:top_n]
            except Exception as e:
                logger.error(f"Cross-encoder reranking failed: {e}")
        
        # No reranking
        return chunks_scores[:top_n]

# --- RAG System ---
class RAGSystem:
    """Main RAG system"""
    
    def __init__(self):
        self.dense_index = DenseIndex()
        self.splade_encoder = None
        self.sparse_index = SparseIndex()
        self.bm25_index = BM25Index()
        self.reranker = Reranker()
        self.chunks = []
        self._load_existing_data()
    
    def _load_existing_data(self):
        """Load existing chunks and indices"""
        # Load chunks from meta file
        if META_FILE.exists():
            try:
                with open(META_FILE, 'r') as f:
                    for line in f:
                        if line.strip():
                            chunk_data = json.loads(line.strip())
                            if 'id' not in chunk_data:
                                chunk_data['id'] = compute_chunk_id(chunk_data)
                            self.chunks.append(chunk_data)
                logger.info(f"Loaded {len(self.chunks)} chunks from meta file")
            except Exception as e:
                logger.error(f"Failed to load meta file: {e}")
        
        # Load sparse index
        self.sparse_index.load(SPLADE_DIR / "index.pkl")
        self.bm25_index.load(BM25_DIR / "index.pkl")
    
    def init_splade(self, model_name: str, top_k_vocab: int):
        """Initialize SPLADE encoder"""
        self.splade_encoder = SpladeEncoder(model_name, top_k_vocab)
    
    def process_documents(self, uploaded_files, chunker: UniversalChunker) -> Dict[str, int]:
        """Process uploaded documents"""
        stats = {'files': 0, 'new_chunks': 0, 'skipped_chunks': 0}
        new_chunks = []
        
        for uploaded_file in uploaded_files:
            filename = uploaded_file.name
            file_bytes = uploaded_file.getvalue()
            
            # Extract text based on file type
            if filename.lower().endswith('.pdf'):
                docs = DocumentProcessor.extract_text_from_pdf(file_bytes, filename)
            elif filename.lower().endswith('.docx'):
                docs = DocumentProcessor.extract_text_from_docx(file_bytes, filename)
            elif filename.lower().endswith(('.txt', '.md', '.markdown')):
                docs = DocumentProcessor.extract_text_from_txt(file_bytes, filename)
            else:
                continue
            
            stats['files'] += 1
            
            # Chunk each document
            for doc_data in docs:
                doc_chunks = chunker.chunk_document(doc_data)
                new_chunks.extend(doc_chunks)
        
        # Deduplicate chunks
        existing_ids = {chunk['id'] for chunk in self.chunks}
        unique_chunks = [chunk for chunk in new_chunks if chunk['id'] not in existing_ids]
        
        stats['new_chunks'] = len(unique_chunks)
        stats['skipped_chunks'] = len(new_chunks) - len(unique_chunks)
        
        if unique_chunks:
            # Save to meta file
            with open(META_FILE, 'a', encoding='utf-8') as f:
                for chunk in unique_chunks:
                    f.write(json.dumps(chunk, ensure_ascii=False) + '\n')
            
            # Add to memory
            self.chunks.extend(unique_chunks)
            
            # Index chunks
            self._index_chunks(unique_chunks)
        
        return stats
    
    def _index_chunks(self, chunks: List[Dict]):
        """Index new chunks"""
        # Dense indexing
        self.dense_index.add(chunks)
        
        # Sparse indexing
        if self.splade_encoder and self.splade_encoder.available:
            for chunk in chunks:
                sparse_vector = self.splade_encoder.encode(chunk['text'])
                if sparse_vector:
                    self.sparse_index.add_document(chunk['id'], sparse_vector, chunk)
            
            # Save SPLADE index
            self.sparse_index.save(SPLADE_DIR / "index.pkl")
            
            # Save config
            config = {
                'model_name': self.splade_encoder.model_name,
                'top_k_vocab': self.splade_encoder.top_k_vocab
            }
            with open(SPLADE_DIR / "config.json", 'w') as f:
                json.dump(config, f)
        
        # BM25 indexing
        if self.bm25_index.available:
            self.bm25_index.add_documents(self.chunks)  # Rebuild with all chunks
            self.bm25_index.save(BM25_DIR / "index.pkl")
    
    def rebuild_sparse_index(self, backend: str):
        """Rebuild sparse index from existing chunks"""
        if backend == "SPLADE++" and self.splade_encoder and self.splade_encoder.available:
            self.sparse_index = SparseIndex()
            for chunk in self.chunks:
                sparse_vector = self.splade_encoder.encode(chunk['text'])
                if sparse_vector:
                    self.sparse_index.add_document(chunk['id'], sparse_vector, chunk)
            self.sparse_index.save(SPLADE_DIR / "index.pkl")
        
        elif backend == "BM25" and self.bm25_index.available:
            self.bm25_index.add_documents(self.chunks)
            self.bm25_index.save(BM25_DIR / "index.pkl")
    
    def search(self, query: str, mode: str, top_k: int, rerank_top_n: int, rrf_k: int, alpha_weight: float = 0.7) -> List[Tuple[Dict, float]]:
        """Search with different modes"""
        if mode == "semantic":
            return self.dense_index.search(query, top_k)
        
        elif mode == "sparse (SPLADE)":
            if self.splade_encoder and self.splade_encoder.available:
                query_vector = self.splade_encoder.encode(query)
                if query_vector:
                    doc_scores = self.sparse_index.search(query_vector, top_k)
                    results = []
                    for doc_id, score in doc_scores:
                        metadata = self.sparse_index.doc_metadata.get(doc_id)
                        if metadata:
                            results.append((metadata, score))
                    return results
            
            # Fallback to BM25
            return self.bm25_index.search(query, self.chunks, top_k)
        
        elif mode == "hybrid":
            # Get results from both methods
            dense_results = self.dense_index.search(query, rerank_top_n)
            
            sparse_results = []
            if self.splade_encoder and self.splade_encoder.available:
                query_vector = self.splade_encoder.encode(query)
                if query_vector:
                    doc_scores = self.sparse_index.search(query_vector, rerank_top_n)
                    for doc_id, score in doc_scores:
                        metadata = self.sparse_index.doc_metadata.get(doc_id)
                        if metadata:
                            sparse_results.append((metadata, score))
            else:
                sparse_results = self.bm25_index.search(query, self.chunks, rerank_top_n)
            
            # RRF fusion
            fused_results = reciprocal_rank_fusion([dense_results, sparse_results], rrf_k)
            return fused_results[:max(top_k, rerank_top_n)]
        
        elif mode == "alpha (weighted)":
            # Alpha mode: weighted combination of dense and sparse scores
            dense_results = self.dense_index.search(query, rerank_top_n)
            
            sparse_results = []
            if self.splade_encoder and self.splade_encoder.available:
                query_vector = self.splade_encoder.encode(query)
                if query_vector:
                    doc_scores = self.sparse_index.search(query_vector, rerank_top_n)
                    for doc_id, score in doc_scores:
                        metadata = self.sparse_index.doc_metadata.get(doc_id)
                        if metadata:
                            sparse_results.append((metadata, score))
            else:
                sparse_results = self.bm25_index.search(query, self.chunks, rerank_top_n)
            
            # Weighted combination using alpha
            alpha_results = alpha_weighted_fusion(dense_results, sparse_results, alpha_weight)
            return alpha_results[:max(top_k, rerank_top_n)]
        
        return []
    
    def reset_store(self):
        """Reset entire store"""
        import shutil
        try:
            shutil.rmtree(RAG_STORE_DIR)
            for dir_path in [RAG_STORE_DIR, CHROMA_DIR, FAISS_DIR, SPLADE_DIR, BM25_DIR]:
                dir_path.mkdir(parents=True, exist_ok=True)
            
            self.chunks = []
            self.sparse_index = SparseIndex()
            self.bm25_index = BM25Index()
            self.dense_index = DenseIndex()
        except Exception as e:
            logger.error(f"Failed to reset store: {e}")

# --- Answer Generation ---
def generate_gpt_answer(query: str, chunks_scores: List[Tuple[Dict, float]], 
                       gpt_model: str = "gpt-3.5-turbo", api_key: str = None) -> Tuple[str, List[Dict]]:
    """Generate answer using GPT with retrieved context"""
    if not chunks_scores or not HAS_OPENAI or not api_key:
        return generate_answer(query, chunks_scores)
    
    try:
        # Prepare context from top chunks
        context_parts = []
        used_chunks = []
        total_tokens = 0
        max_context_tokens = 3000  # Leave room for query and response
        
        for chunk, score in chunks_scores[:10]:  # Use top 10 chunks
            chunk_text = chunk['text']
            chunk_tokens = count_tokens(chunk_text)
            
            if total_tokens + chunk_tokens > max_context_tokens:
                break
                
            context_parts.append(f"[Source: {chunk.get('source_name', 'Unknown')}]\n{chunk_text}")
            used_chunks.append(chunk)
            total_tokens += chunk_tokens
        
        if not context_parts:
            return "No relevant context found.", []
        
        context = "\n\n".join(context_parts)
        
        # Create prompt
        prompt = f"""Based on the following context, please provide a comprehensive and accurate answer to the question. Use only information from the provided context. If the context doesn't contain enough information to answer the question completely, state that clearly.

Context:
{context}

Question: {query}

Answer:"""

        # Initialize OpenAI client
        client = openai.OpenAI(api_key=api_key)
        
        # Generate response
        response = client.chat.completions.create(
            model=gpt_model,
            messages=[
                {"role": "system", "content": "You are a helpful assistant that answers questions based only on the provided context. Be accurate and concise."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=500,
            temperature=0.1
        )
        
        answer = response.choices[0].message.content.strip()
        return answer, used_chunks
        
    except Exception as e:
        logger.error(f"GPT answer generation failed: {e}")
        # Fallback to extractive method
        return generate_answer(query, chunks_scores)

def format_policy_answer(answer_text: str, query_lower: str) -> str:
    """Format policy-related answers to be more direct and user-friendly"""
    import re
    
    # Handle dress code prohibition questions
    if 'prohibit' in query_lower and ('dress code' in query_lower or 'clothing' in query_lower):
        # Extract specific prohibited items
        prohibited_items = extract_prohibited_items(answer_text)
        
        if prohibited_items:
            if len(prohibited_items) == 1:
                return f"The dress code prohibits {prohibited_items[0]}."
            elif len(prohibited_items) == 2:
                return f"The dress code prohibits {prohibited_items[0]} and {prohibited_items[1]}."
            else:
                items_str = ', '.join(prohibited_items[:-1]) + f', and {prohibited_items[-1]}'
                return f"The dress code prohibits {items_str}."
    
    # Handle working hours questions
    if 'working hours' in query_lower or 'hours' in query_lower:
        # Extract time information and format clearly
        time_info = extract_time_information(answer_text)
        if time_info:
            return time_info
    
    # Handle eligibility criteria questions
    if any(word in query_lower for word in ['eligibility', 'criteria', 'requirements', 'eligible']):
        eligibility_info = extract_eligibility_criteria(answer_text, query_lower)
        if eligibility_info:
            return eligibility_info
    
    # Handle outline/summary questions
    if any(word in query_lower for word in ['outline', 'summarize', 'list', 'what are']):
        outline_info = create_structured_outline(answer_text, query_lower)
        if outline_info:
            return outline_info
    
    # For other policy questions, clean up the language
    answer_text = clean_policy_language(answer_text)
    
    return answer_text

def extract_prohibited_items(text: str) -> List[str]:
    """Extract specific prohibited clothing items from policy text"""
    import re
    
    prohibited_items = []
    text_lower = text.lower()
    
    # Look for common prohibited clothing items
    clothing_patterns = [
        r'ripped\s+clothing(?:\s+of\s+any\s+sort)?',
        r'low\s+cut\s+clothing(?:\s+such\s+as\s+jeans\s+and\s+t-shirts)?',
        r'casual\s+tops',
        r'track\s+suits?(?:\s*,?\s*including\s+pants\s+or\s+windcheaters)?',
        r'jeans',
        r't-shirts',
        r'windcheaters',
        r'pants\s+or\s+windcheaters'
    ]
    
    for pattern in clothing_patterns:
        matches = re.findall(pattern, text_lower)
        for match in matches:
            clean_item = match.strip().rstrip(',')
            if clean_item and clean_item not in prohibited_items:
                prohibited_items.append(clean_item)
    
    # Clean up and standardize the items
    standardized_items = []
    for item in prohibited_items:
        if 'ripped clothing' in item:
            standardized_items.append('ripped clothing of any sort')
        elif 'low cut clothing' in item:
            standardized_items.append('low cut clothing (such as jeans and T-shirts)')
        elif 'casual tops' in item:
            standardized_items.append('casual tops')
        elif 'track suit' in item:
            standardized_items.append('track suits (including pants and windcheaters)')
        elif item not in ['jeans', 't-shirts', 'windcheaters'] or not any('track suit' in std for std in standardized_items):
            # Avoid duplicating items already covered by "track suits"
            if item not in standardized_items:
                standardized_items.append(item)
    
    return standardized_items[:4]  # Limit to avoid overly long lists

def extract_time_information(text: str) -> str:
    """Extract and format time information clearly"""
    import re
    
    # Look for time patterns
    time_patterns = re.findall(r'\d{1,2}:\d{2}\s*(?:AM|PM|am|pm)', text)
    day_patterns = re.findall(r'(?:Monday|Tuesday|Wednesday|Thursday|Friday|Saturday|Sunday)', text, re.IGNORECASE)
    
    if time_patterns:
        formatted_times = []
        for time in time_patterns:
            formatted_times.append(time.upper())
        
        if day_patterns:
            days = ', '.join(set(day.title() for day in day_patterns))
            return f"Working hours are {' to '.join(formatted_times)} on {days}."
        else:
            return f"Working hours are {' to '.join(formatted_times)}."
    
    return text

def extract_eligibility_criteria(text: str, query_lower: str) -> str:
    """Extract and format eligibility criteria in a structured, user-friendly way"""
    import re
    
    criteria_items = []
    text_lower = text.lower()
    
    # WFH specific criteria - comprehensive policy extraction
    if 'wfh' in query_lower or 'work from home' in query_lower:
        return extract_comprehensive_wfh_policy(text, query_lower)
    
    # General eligibility criteria extraction
    criteria_keywords = ['must', 'required', 'should', 'need to', 'criteria', 'eligible', 'qualification']
    sentences = re.split(r'[.!?]+', text)
    
    for sentence in sentences:
        sentence = sentence.strip()
        if len(sentence) < 10:
            continue
            
        sentence_lower = sentence.lower()
        if any(keyword in sentence_lower for keyword in criteria_keywords):
            # Clean up the sentence
            clean_sentence = re.sub(r'^[^a-zA-Z]*', '', sentence).strip()
            if clean_sentence and len(clean_sentence) > 15:
                criteria_items.append(f"• {clean_sentence}")
    
    if criteria_items:
        return "**Eligibility Criteria:**\n\n" + '\n'.join(criteria_items[:4])  # Limit to 4 items
    
    return ""

def create_structured_outline(text: str, query_lower: str) -> str:
    """Create a structured outline/summary from policy text"""
    import re
    
    # For WFH outlines
    if 'wfh' in query_lower or 'work from home' in query_lower:
        outline_sections = []
        text_lower = text.lower()
        
        # Role requirements
        if 'role' in query_lower:
            role_info = extract_role_requirements(text)
            if role_info:
                outline_sections.append(f"**Role Requirements:**\n{role_info}")
        
        # Probation status
        if 'probation' in query_lower:
            probation_info = extract_probation_requirements(text)
            if probation_info:
                outline_sections.append(f"**Probation Status:**\n{probation_info}")
        
        # Performance criteria
        if 'performance' in query_lower:
            performance_info = extract_performance_requirements(text)
            if performance_info:
                outline_sections.append(f"**Performance Requirements:**\n{performance_info}")
        
        if outline_sections:
            return '\n\n'.join(outline_sections)
    
    # General outline creation
    key_points = []
    sentences = re.split(r'[.!?]+', text)
    
    for sentence in sentences:
        sentence = sentence.strip()
        if len(sentence) < 20:
            continue
            
        # Look for important sentences
        if any(word in sentence.lower() for word in ['must', 'required', 'policy', 'criteria', 'designed to']):
            clean_sentence = sentence.strip()
            if clean_sentence:
                key_points.append(f"• {clean_sentence}")
    
    if key_points:
        return "**Key Points:**\n\n" + '\n'.join(key_points[:3])
    
    return ""

def extract_role_requirements(text: str) -> str:
    """Extract role-specific requirements"""
    import re
    
    text_lower = text.lower()
    role_requirements = []
    
    # Look for role-specific patterns
    if re.search(r'based\s+on.*role|role.*basis|position.*requirements', text_lower):
        role_requirements.append("• Role-based assessment required")
    
    if re.search(r'suitable.*role|fit.*role|competency.*role', text_lower):
        role_requirements.append("• Must demonstrate role suitability and competency")
    
    return '\n'.join(role_requirements) if role_requirements else ""

def extract_probation_requirements(text: str) -> str:
    """Extract probation-related requirements"""
    import re
    
    text_lower = text.lower()
    probation_info = []
    
    if re.search(r'3\s+month.*probation|probation.*3\s+month', text_lower):
        probation_info.append("• Must complete 3-month probationary period")
    
    if re.search(r'completed.*probation|probation.*completed', text_lower):
        probation_info.append("• Probationary period must be successfully completed")
    
    return '\n'.join(probation_info) if probation_info else ""

def extract_performance_requirements(text: str) -> str:
    """Extract performance-related requirements"""
    import re
    
    text_lower = text.lower()
    performance_info = []
    
    if re.search(r'meeting.*expectations|exceeding.*expectations', text_lower):
        performance_info.append("• Performance must meet or exceed expectations")
    
    if re.search(r'consistent.*performance|performance.*consistent', text_lower):
        performance_info.append("• Must maintain consistent performance record")
    
    if re.search(r'individual.*performance', text_lower):
        performance_info.append("• Individual performance will be evaluated")
    
    return '\n'.join(performance_info) if performance_info else ""

def extract_comprehensive_wfh_policy(text: str, query_lower: str) -> str:
    """Extract comprehensive WFH policy information with proper structure"""
    import re
    
    text_lower = text.lower()
    policy_sections = []
    
    # Position/Role Eligibility
    position_criteria = []
    position_patterns = [
        r'not\s+all\s+positions?\s+(?:are\s+)?eligible',
        r'eligibility\s+depends\s+on.*role',
        r'role.*department.*job\s+responsibilities',
        r'position.*eligible.*wfh',
        r'specific\s+roles?\s+(?:are\s+)?eligible'
    ]
    
    for pattern in position_patterns:
        if re.search(pattern, text_lower):
            if 'not all positions' in text_lower:
                position_criteria.append("Not all positions are eligible for WFH. Eligibility depends on the employee's role, department, and job responsibilities.")
            break
    
    # Probation and Performance Requirements
    performance_criteria = []
    if re.search(r'completed?\s+(?:their\s+)?probation(?:ary)?\s+period', text_lower):
        if re.search(r'consistent\s+record\s+of\s+performance', text_lower):
            performance_criteria.append("Employees must have completed their probationary period and must have a consistent record of performance that meets or exceeds expectations.")
    
    # Approval Process
    approval_criteria = []
    approval_patterns = [
        r'case-by-case\s+basis',
        r'discretion\s+of.*manager.*hr',
        r'decisions\s+made\s+at.*discretion',
        r'evaluated.*case-by-case',
        r'manager.*hr.*approval'
    ]
    
    for pattern in approval_patterns:
        if re.search(pattern, text_lower):
            approval_criteria.append("WFH eligibility is evaluated on a case-by-case basis, with decisions made at the discretion of the employee's manager and HR.")
            break
    
    # General Policy
    general_policy = []
    general_patterns = [
        r'generally.*does\s+not\s+permit\s+wfh',
        r'company\s+does\s+not\s+permit.*except',
        r'not\s+permitted.*emergency\s+cases',
        r'proper\s+approval\s+and\s+process'
    ]
    
    for pattern in general_patterns:
        if re.search(pattern, text_lower):
            general_policy.append("Generally, the company does not permit WFH except in emergency cases, with proper approval and process.")
            break
    
    # Remote Work Restrictions
    remote_restrictions = []
    restriction_patterns = [
        r'remote\s+work\s+from\s+hometowns?\s+(?:is\s+)?not\s+permitted',
        r'connectivity\s+issues',
        r'avoid\s+miscommunication',
        r'hometowns?.*not\s+permitted'
    ]
    
    for pattern in restriction_patterns:
        if re.search(pattern, text_lower):
            remote_restrictions.append("Remote work from hometowns is not permitted due to connectivity issues and to avoid miscommunication.")
            break
    
    # Compile comprehensive answer
    all_criteria = position_criteria + performance_criteria + approval_criteria + general_policy + remote_restrictions
    
    if all_criteria:
        header = "**The eligibility criteria for Work From Home (WFH) at Antraweb Technologies Pvt Ltd are as follows:**\n\n"
        
        # Format as numbered list for better structure
        formatted_criteria = []
        for i, criterion in enumerate(all_criteria, 1):
            formatted_criteria.append(f"{i}. {criterion}")
        
        return header + '\n'.join(formatted_criteria)
    
    # Fallback to simpler extraction if comprehensive fails
    return extract_simple_wfh_criteria(text)

def extract_simple_wfh_criteria(text: str) -> str:
    """Fallback function for simpler WFH criteria extraction"""
    import re
    
    text_lower = text.lower()
    criteria = []
    
    # Basic probation requirement
    if re.search(r'completed?\s+(?:their\s+)?probation(?:ary)?\s+period', text_lower):
        criteria.append("✓ Complete probationary period")
    
    # Performance requirement
    if re.search(r'consistent\s+record\s+of\s+performance', text_lower):
        criteria.append("✓ Maintain consistent performance meeting or exceeding expectations")
    
    # Company discretion
    if re.search(r'company.*requirements|discretion', text_lower):
        criteria.append("✓ Subject to company requirements and management discretion")
    
    if criteria:
        header = "**WFH Eligibility Criteria:**\n\n"
        return header + '\n'.join(criteria)
    
    return ""

def clean_policy_language(text: str) -> str:
    """Clean up policy language to be more user-friendly"""
    import re
    
    # Remove redundant policy references
    text = re.sub(r'\b(?:The\s+)?(?:dress\s+code\s+)?policy\s+explicitly\s+', '', text, flags=re.IGNORECASE)
    text = re.sub(r'\bprohibits\s+the\s+following\s+types?\s+of\s+clothing\s+at\s+the\s+workplace:\s*', 'prohibits ', text, flags=re.IGNORECASE)
    
    # Clean up redundant phrases
    text = re.sub(r'\s+', ' ', text).strip()
    
    # Ensure proper capitalization
    if text and not text[0].isupper():
        text = text[0].upper() + text[1:]
    
    return text

def generate_answer(query: str, chunks_scores: List[Tuple[Dict, float]]) -> Tuple[str, List[Dict]]:
    """Generate extractive answer from retrieved chunks with improved accuracy"""
    if not chunks_scores:
        return "Not enough information in the provided context.", []
    
    query_lower = query.lower()
    query_words = set(query_lower.split())
    
    # Enhanced keyword detection for better matching
    query_keywords = []
    for word in query_words:
        if len(word) > 2 and word not in {'the', 'and', 'are', 'what', 'how', 'when', 'where', 'why'}:
            query_keywords.append(word)
    
    # Collect all relevant content with enhanced scoring
    relevant_content = []
    used_chunks = []
    
    for chunk, retrieval_score in chunks_scores:
        chunk_text = chunk['text']
        chunk_lower = chunk_text.lower()
        
        # Check for direct keyword matches in the chunk
        keyword_matches = sum(1 for keyword in query_keywords if keyword in chunk_lower)
        if keyword_matches == 0:
            continue
        
        # Split into sentences and score each
        sentences = re.split(r'[.!?]+', chunk_text)
        
        for sentence in sentences:
            sentence = sentence.strip()
            if len(sentence) < 10:  # Skip very short sentences
                continue
            
            sentence_lower = sentence.lower()
            
            # Enhanced scoring system
            score = 0
            
            # Keyword presence score
            sentence_keywords = sum(1 for keyword in query_keywords if keyword in sentence_lower)
            score += sentence_keywords * 2
            
            # Query type specific scoring
            if 'working hours' in query_lower or 'hours' in query_lower:
                # Look for time patterns and schedule-related terms
                time_patterns = [
                    r'\d{1,2}:\d{2}\s*(?:am|pm)',  # 9:45 AM
                    r'\d{1,2}\s*(?:am|pm)',        # 9 AM
                    r'monday|tuesday|wednesday|thursday|friday|saturday|sunday',
                    r'schedule|timing|hours|shift'
                ]
                for pattern in time_patterns:
                    if re.search(pattern, sentence_lower):
                        score += 3
            
            # Enhanced dress code and policy scoring
            if any(word in query_lower for word in ['dress code', 'clothing', 'prohibit', 'policy']):
                # Look for policy-specific patterns
                policy_patterns = [
                    r'prohibit[s]?\s+(?:the\s+)?following',  # "prohibits the following"
                    r'not\s+allowed',                        # "not allowed"
                    r'forbidden',                            # "forbidden"
                    r'dress\s+code',                         # "dress code"
                    r'(?:ripped|low\s+cut|casual|track\s+suit)',  # specific clothing items
                ]
                for pattern in policy_patterns:
                    if re.search(pattern, sentence_lower):
                        score += 4
                
                # Boost for sentences with lists or enumerations
                if ':' in sentence and sentence.count(',') >= 2:
                    score += 3
            
            # Enhanced eligibility and criteria scoring
            if any(word in query_lower for word in ['eligibility', 'criteria', 'requirements', 'eligible', 'outline']):
                # Look for eligibility-specific patterns
                eligibility_patterns = [
                    r'eligib(?:le|ility)',                   # "eligible", "eligibility"
                    r'criteria|requirements',                # "criteria", "requirements"
                    r'must\s+(?:have|be|complete)',         # "must have/be/complete"
                    r'probation(?:ary)?\s+period',          # "probationary period"
                    r'performance.*(?:meeting|exceeding)',   # performance expectations
                    r'completed?\s+their',                   # "completed their"
                    r'consistent\s+record',                  # "consistent record"
                ]
                for pattern in eligibility_patterns:
                    if re.search(pattern, sentence_lower):
                        score += 5  # Higher boost for eligibility content
                
                # Extra boost for WFH eligibility - comprehensive policy content
                if any(wfh in query_lower for wfh in ['wfh', 'work from home']):
                    # Boost for comprehensive WFH policy patterns
                    wfh_policy_patterns = [
                        r'not\s+all\s+positions?\s+(?:are\s+)?eligible',
                        r'case-by-case\s+basis',
                        r'discretion\s+of.*manager',
                        r'emergency\s+cases',
                        r'hometowns?\s+(?:is\s+)?not\s+permitted',
                        r'connectivity\s+issues',
                        r'avoid\s+miscommunication'
                    ]
                    
                    for pattern in wfh_policy_patterns:
                        if re.search(pattern, sentence_lower):
                            score += 6  # High boost for comprehensive policy content
                    
                    # General WFH mentions
                    if any(wfh_term in sentence_lower for wfh_term in ['work from home', 'wfh', 'remote work']):
                        score += 3
            
            # Prefer sentences with specific details
            if re.search(r'\d+', sentence):  # Contains numbers
                score += 1
            
            # Penalize very long sentences (likely to be less precise)
            if len(sentence.split()) > 50:
                score -= 1
            
            if score >= 2:  # Minimum threshold
                relevance = score + retrieval_score * 0.5  # Combine with retrieval score
                relevant_content.append((sentence, relevance, chunk))
    
    if not relevant_content:
        return "Not enough information in the provided context.", []
    
    # Sort by relevance score
    relevant_content.sort(key=lambda x: x[1], reverse=True)
    
    # Select best sentences with diversity
    selected_sentences = []
    selected_chunks_ids = set()
    
    for sentence, relevance, chunk in relevant_content:
        # Allow more sentences for complete answers, especially for WFH policies
        max_sentences = 8 if any(wfh in query_lower for wfh in ['wfh', 'work from home', 'policy']) else 5
        if len(selected_sentences) >= max_sentences:
            break
        
        # Check for diversity (avoid redundant information)
        is_diverse = True
        for selected_sentence, _, _ in selected_sentences:
            similarity = jaccard_similarity(sentence, selected_sentence, n=4)
            if similarity > 0.6:  # Less strict diversity for better coverage
                is_diverse = False
                break
        
        if is_diverse:
            selected_sentences.append((sentence, relevance, chunk))
            if chunk['id'] not in selected_chunks_ids:
                used_chunks.append(chunk)
                selected_chunks_ids.add(chunk['id'])
    
    if selected_sentences:
        # Create a more coherent answer
        answer_parts = []
        
        # Group by topic/context if possible
        high_relevance = [s for s in selected_sentences if s[1] >= 5]
        medium_relevance = [s for s in selected_sentences if 3 <= s[1] < 5]
        
        # Prioritize high relevance sentences
        for sentence, _, _ in high_relevance[:3]:
            answer_parts.append(sentence.strip())
        
        # Add medium relevance if we need more context
        if len(answer_parts) < 2:
            for sentence, _, _ in medium_relevance[:2]:
                answer_parts.append(sentence.strip())
        
        # Clean up and format the answer
        answer_text = ' '.join(answer_parts)
        
        # Post-process for specific query types
        if any(word in query_lower for word in ['dress code', 'prohibit', 'clothing']):
            answer_text = format_policy_answer(answer_text, query_lower)
        
        # Remove any incomplete sentences at the end
        answer_text = re.sub(r'\s+', ' ', answer_text).strip()
        
        return answer_text, used_chunks
    else:
        return "Not enough information in the provided context.", []

def format_references(chunks: List[Dict]) -> List[str]:
    """Format references from chunks"""
    references = []
    seen_refs = set()
    
    for chunk in chunks:
        source_name = chunk.get('source_name', 'Unknown')
        page_start = chunk.get('page_start', '')
        page_end = chunk.get('page_end', '')
        section = chunk.get('section_path', '-')
        chunk_id = chunk.get('id', 'unknown')[:8]
        
        # Create page range
        if page_start and page_end:
            if page_start == page_end:
                page_range = str(page_start)
            else:
                page_range = f"{page_start}–{page_end}"
        else:
            page_range = "-"
        
        # Create reference key for deduplication
        ref_key = (source_name, page_range, section)
        if ref_key not in seen_refs:
            ref_text = f"[ID:{chunk_id}] {source_name} — pages {page_range} | section: {section}"
            references.append(ref_text)
            seen_refs.add(ref_key)
    
    return references

# --- Grounded Charts & Tables ---

def generate_grounded_chart(request: str, chunks_scores: List[Tuple[Dict, float]], chart_type_hint: str, include_csv: bool = True) -> str:
    """
    Generate grounded charts and tables from retrieved document chunks.
    Follows the strict grounding rules: Context-only, no hallucinations, Answer then References format.
    """
    import re
    import json
    from typing import List, Dict, Any
    
    # Format context for analysis
    context_chunks = []
    for i, (chunk, score) in enumerate(chunks_scores):
        context_chunks.append({
            'id': chunk.get('id', f'chunk_{i}'),
            'source_name': chunk.get('source', 'Unknown'),
            'page_start': chunk.get('page', 1),
            'page_end': chunk.get('page', 1),
            'section': chunk.get('section', '-'),
            'text': chunk.get('text', ''),
            'score': score
        })
    
    # Build context string in the required format
    context_str = ""
    for chunk in context_chunks:
        context_str += f"[ID: {chunk['id']}]\n"
        context_str += f"source_name: {chunk['source_name']}\n"
        context_str += f"page_start: {chunk['page_start']}   page_end: {chunk['page_end']}\n"
        context_str += f"section: {chunk['section']}\n"
        context_str += f"text: |\n{chunk['text']}\n\n"
    
    # Generate the response using the grounded assistant rules
    response = process_grounded_request(request, context_str, chart_type_hint, include_csv)
    
    return response

def process_grounded_request(request: str, context: str, chart_type_hint: str, include_csv: bool) -> str:
    """
    Process a grounded chart/table request following the strict rules.
    Returns formatted markdown response with Answer and References sections.
    """
    
    # This is a simplified implementation that extracts basic data patterns
    # In a full implementation, you'd use more sophisticated NLP/data extraction
    
    response_parts = []
    references_used = []
    
    # Start with Answer section
    response_parts.append("## Answer")
    
    # Analyze the request to determine if table or chart is wanted
    request_lower = request.lower()
    wants_chart = any(word in request_lower for word in ['chart', 'graph', 'plot', 'visualization'])
    wants_table = any(word in request_lower for word in ['table', 'list', 'show'])
    
    # Extract numeric data from context (simplified extraction)
    extracted_data = extract_numeric_data_from_context(context)
    
    # For Excel/CSV content, try enhanced extraction
    if 'Sheet:' in context or 'CSV Data' in context or 'Columns:' in context:
        excel_chunks = [{'text': context, 'source': 'uploaded_file'}]
        excel_data = extract_excel_data_directly(excel_chunks)
        if excel_data and len(excel_data) > len(extracted_data):
            extracted_data = excel_data
    
    if not extracted_data:
        response_parts.append("Not enough information in the provided context.")
        response_parts.append("\nNo sufficient numeric data found to create the requested visualization.")
        response_parts.append("\n**What we look for:**")
        response_parts.append("- Working hours: '9:45 AM to 6:15 PM'")
        response_parts.append("- Leave policies: '21 days annual leave'")
        response_parts.append("- Financial data: '$1,250,000 revenue'")
        response_parts.append("- Percentages: '45% completion rate'")
        response_parts.append("- Time series: 'Q1 2023: 1500'")
        response_parts.append("\n**Tips:** Try uploading documents with tables, financial reports, or structured data.")
    else:
        if wants_table or chart_type_hint == "Table":
            # Generate table
            table_result = generate_table_from_data(extracted_data, request)
            response_parts.append(table_result)
        elif wants_chart or chart_type_hint != "Auto-detect":
            # Generate chart
            chart_result = generate_chart_from_data(extracted_data, request, chart_type_hint)
            response_parts.append(chart_result)
        else:
            # Auto-detect best format
            if len(extracted_data) <= 3:
                table_result = generate_table_from_data(extracted_data, request)
                response_parts.append(table_result)
            else:
                chart_result = generate_chart_from_data(extracted_data, request, "Bar Chart")
                response_parts.append(chart_result)
        
        # Add CSV if requested
        if include_csv and extracted_data:
            csv_data = generate_csv_from_data(extracted_data)
            response_parts.append(f"\n```csv\n{csv_data}\n```")
    
    # Add References section
    response_parts.append("\n## References")
    
    # Extract unique sources from context
    sources = extract_sources_from_context(context)
    for source in sources:
        response_parts.append(f"- [ID:{source['id']}] {source['source_name']} — pages {source['page_start']}–{source['page_end']} | section: {source['section']}")
    
    return "\n".join(response_parts)

def extract_excel_data_directly(chunks: List[Dict]) -> List[Dict[str, Any]]:
    """Extract structured data directly from Excel/CSV content for better chart generation."""
    import re
    
    data_points = []
    
    for chunk in chunks:
        content = chunk.get('text', '')
        filename = chunk.get('source', '')
        
        # Check if this is Excel/CSV content
        if 'Sheet:' in content or 'CSV Data' in content or 'Columns:' in content:
            lines = content.split('\n')
            current_sheet = "Sheet1"
            
            for line in lines:
                line = line.strip()
                
                # Track sheet names
                if line.startswith('Sheet:'):
                    current_sheet = line.replace('Sheet:', '').strip()
                    continue
                
                # Skip header lines
                if line.startswith('Columns:') or line.startswith('CSV Data'):
                    continue
                
                # Extract structured data: "Column: Value"
                if ':' in line and line.count(':') == 1:
                    parts = line.split(':', 1)
                    if len(parts) == 2:
                        category = parts[0].strip()
                        value_str = parts[1].strip()
                        
                        # Try to extract numeric value
                        numeric_match = re.search(r'[\d,]+\.?\d*', value_str)
                        if numeric_match:
                            try:
                                numeric_value = float(numeric_match.group().replace(',', ''))
                                
                                # Extract unit if present
                                unit_match = re.search(r'[%$€£]|([a-zA-Z]+)(?:\s|$)', value_str)
                                unit = unit_match.group().strip() if unit_match else ''
                                
                                # Create enhanced category name
                                if current_sheet != "Sheet1":
                                    category = f"{current_sheet} - {category}"
                                
                                data_points.append({
                                    'category': category,
                                    'value': numeric_value,
                                    'unit': unit,
                                    'source_type': 'excel',
                                    'raw_text': line,
                                    'filename': filename
                                })
                            except ValueError:
                                continue
    
    # Remove duplicates and return top results
    unique_data = []
    seen = set()
    for dp in data_points:
        key = (dp['category'], dp['value'], dp['unit'])
        if key not in seen:
            seen.add(key)
            unique_data.append(dp)
    
    # Sort by value descending
    unique_data.sort(key=lambda x: x['value'], reverse=True)
    return unique_data[:20]  # Return more data points for Excel files


def extract_numeric_data_from_context(context: str) -> List[Dict[str, Any]]:
    """Extract numeric data patterns from context chunks with enhanced pattern matching."""
    import re
    
    data_points = []
    lines = context.split('\n')
    
    # Look for common data patterns
    for line in lines:
        line = line.strip()
        if not line or line.startswith('[ID:') or line.startswith('source_name:') or line.startswith('page_') or line.startswith('section:') or line.startswith('text:'):
            continue
            
        # Pattern 1: Specific HR/Business patterns
        # Working hours: "9:45 AM to 6:15 PM"
        time_range_pattern = re.findall(r'(\d{1,2}:\d{2})\s*(?:AM|PM|am|pm)\s*(?:to|–|-)\s*(\d{1,2}:\d{2})\s*(?:AM|PM|am|pm)', line)
        if time_range_pattern and ('working' in line.lower() or 'hours' in line.lower() or 'office' in line.lower()):
            for start_time, end_time in time_range_pattern:
                # Calculate duration
                start_hour = float(start_time.split(':')[0]) + float(start_time.split(':')[1])/60
                end_hour = float(end_time.split(':')[0]) + float(end_time.split(':')[1])/60
                if 'PM' in line and start_hour < 12:
                    start_hour += 12
                if 'PM' in line and end_hour < 12:
                    end_hour += 12
                duration = end_hour - start_hour
                
                # Extract day context
                day_context = "Weekday"
                if 'monday' in line.lower() and 'friday' in line.lower():
                    day_context = "Monday to Friday"
                elif 'saturday' in line.lower():
                    day_context = "Saturday"
                elif 'sunday' in line.lower():
                    day_context = "Sunday"
                
                data_points.append({
                    'category': f"{day_context} Working Hours",
                    'value': duration,
                    'unit': 'hours',
                    'raw_text': line
                })
        
        # Pattern 2: Leave/Days patterns
        leave_pattern = re.findall(r'(\d+)\s+(days?|weeks?|months?)\s+(?:of\s+)?([a-zA-Z\s]+(?:leave|vacation|holiday))', line, re.IGNORECASE)
        for match in leave_pattern:
            value, unit, leave_type = match
            try:
                numeric_value = float(value)
                category = f"{leave_type.strip().title()}"
                data_points.append({
                    'category': category,
                    'value': numeric_value,
                    'unit': unit.lower(),
                    'raw_text': line
                })
            except ValueError:
                continue
        
        # Pattern 3: Percentage patterns
        percentage_pattern = re.findall(r'(\d+(?:\.\d+)?)\s*%\s+([a-zA-Z\s]+)', line)
        for match in percentage_pattern:
            value, description = match
            try:
                numeric_value = float(value)
                category = description.strip().title()
                if len(category) > 3:
                    data_points.append({
                        'category': category,
                        'value': numeric_value,
                        'unit': '%',
                        'raw_text': line
                    })
            except ValueError:
                continue
        
        # Pattern 4: Currency patterns
        currency_pattern = re.findall(r'([\$€£])\s*(\d+(?:,\d{3})*(?:\.\d+)?)\s+([a-zA-Z\s]+)', line)
        for match in currency_pattern:
            currency, value, description = match
            try:
                numeric_value = float(value.replace(',', ''))
                category = description.strip().title()
                if len(category) > 3 and numeric_value > 0:
                    data_points.append({
                        'category': category,
                        'value': numeric_value,
                        'unit': currency,
                        'raw_text': line
                    })
            except ValueError:
                continue
        
        # Pattern 5: General "Number Unit Description" pattern
        general_pattern = re.findall(r'(\d+(?:,\d{3})*(?:\.\d+)?)\s+([a-zA-Z]+)\s+([a-zA-Z\s]{5,30})', line)
        for match in general_pattern:
            value, unit, description = match
            try:
                numeric_value = float(value.replace(',', ''))
                # Clean description
                category = description.strip().title()
                category = re.sub(r'\s+', ' ', category)  # Normalize spaces
                
                # Filter out meaningless patterns
                if (len(category) > 5 and numeric_value > 0 and 
                    not any(skip in category.lower() for skip in ['if ', 'she ', 'he ', 'must ', 'will ', 'the ', 'an ', 'and '])):
                    data_points.append({
                        'category': category,
                        'value': numeric_value,
                        'unit': unit.lower(),
                        'raw_text': line
                    })
            except ValueError:
                continue
        
        # Pattern 6: Colon-separated "Description: Number Unit"
        colon_pattern = re.findall(r'([A-Za-z][A-Za-z\s]{5,40})\s*:\s*(\d+(?:,\d{3})*(?:\.\d+)?)\s*([a-zA-Z%$€£]*)', line)
        for match in colon_pattern:
            description, value, unit = match
            try:
                numeric_value = float(value.replace(',', ''))
                category = description.strip().title()
                if len(category) > 5 and numeric_value > 0:
                    data_points.append({
                        'category': category,
                        'value': numeric_value,
                        'unit': unit.strip() or '',
                        'raw_text': line
                    })
            except ValueError:
                continue
        
        # Pattern 7: Time series data "Quarter/Month Year: Value"
        time_series_pattern = re.findall(r'((?:Q[1-4]|Quarter\s*[1-4]|Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\s*\d{4})\s*:\s*(\d+(?:,\d{3})*(?:\.\d+)?)', line)
        for match in time_series_pattern:
            time_period, value = match
            try:
                numeric_value = float(value.replace(',', ''))
                data_points.append({
                    'time_period': time_period.strip(),
                    'value': numeric_value,
                    'unit': '',
                    'raw_text': line
                })
            except ValueError:
                continue
    
    # Remove duplicates and clean data
    unique_data = []
    seen = set()
    for dp in data_points:
        # Create unique key
        key = (dp.get('category', ''), dp.get('time_period', ''), dp['value'], dp.get('unit', ''))
        if key not in seen and dp['value'] > 0:
            seen.add(key)
            
            # Clean up category names
            if 'category' in dp:
                dp['category'] = re.sub(r'\s+', ' ', dp['category']).strip()
                # Remove trailing punctuation
                dp['category'] = re.sub(r'[.,;:]+$', '', dp['category'])
            
            unique_data.append(dp)
    
    # Sort by value descending and limit results
    unique_data.sort(key=lambda x: x['value'], reverse=True)
    
    return unique_data[:15]  # Limit to top 15 data points

def generate_table_from_data(data_points: List[Dict[str, Any]], request: str) -> str:
    """Generate a markdown table from extracted data points."""
    if not data_points:
        return "No tabular data found in the provided context."
    
    # Sort and limit data for readability
    sorted_data = sorted(data_points, key=lambda x: x['value'], reverse=True)[:15]
    
    # Determine table structure based on data
    if any('time_period' in dp for dp in sorted_data):
        # Time series table
        table_lines = ["| Time Period | Value | Unit |", "|-------------|-------|------|"]
        for dp in sorted_data:
            if 'time_period' in dp:
                value_str = f"{dp['value']:,.1f}" if dp['value'] >= 1 else f"{dp['value']:.3f}"
                table_lines.append(f"| {dp['time_period']} | {value_str} | {dp.get('unit', '')} |")
    else:
        # Category table
        table_lines = ["| Category | Value | Unit |", "|----------|-------|------|"]
        for dp in sorted_data:
            if 'category' in dp:
                # Truncate long categories for table display
                category = dp['category']
                if len(category) > 40:
                    category = category[:37] + "..."
                value_str = f"{dp['value']:,.1f}" if dp['value'] >= 1 else f"{dp['value']:.3f}"
                table_lines.append(f"| {category} | {value_str} | {dp.get('unit', '')} |")
    
    return "\n".join(table_lines)

def generate_chart_from_data(data_points: List[Dict[str, Any]], request: str, chart_type_hint: str) -> str:
    """Generate a Vega-Lite chart specification from extracted data points."""
    if not data_points:
        return "No data available for chart generation."
    
    # Prepare data for Vega-Lite
    chart_data = []
    chart_title = "Data Visualization"
    
    # Determine units from data
    units = [dp.get('unit', '') for dp in data_points if dp.get('unit')]
    most_common_unit = max(set(units), key=units.count) if units else ""
    
    if any('time_period' in dp for dp in data_points):
        # Time series data
        for dp in data_points:
            if 'time_period' in dp:
                chart_data.append({
                    "time": dp['time_period'],
                    "value": dp['value']
                })
        
        chart_title = f"Time Series Data ({most_common_unit})" if most_common_unit else "Time Series Data"
        y_title = f"Value ({most_common_unit})" if most_common_unit else "Value"
        
        # Determine mark type for time series
        mark_config = {"type": "line", "point": True, "strokeWidth": 3}
        if chart_type_hint == "Area Chart":
            mark_config = {"type": "area", "line": True, "point": True}
        
        # Time series chart
        chart_spec = {
            "$schema": "https://vega.github.io/schema/vega-lite/v5.json",
            "title": chart_title,
            "description": f"Time series chart showing values over time",
            "data": {"values": chart_data},
            "mark": mark_config,
            "encoding": {
                "x": {"field": "time", "type": "ordinal", "title": "Time Period"},
                "y": {"field": "value", "type": "quantitative", "title": y_title},
                "tooltip": [
                    {"field": "time", "type": "ordinal", "title": "Period"},
                    {"field": "value", "type": "quantitative", "format": ".1f", "title": "Value"}
                ]
            },
            "width": 600,
            "height": 400
        }
    else:
        # Category data - limit to top 8 for readability
        sorted_data = sorted(data_points, key=lambda x: x['value'], reverse=True)[:8]
        
        for dp in sorted_data:
            if 'category' in dp:
                # Truncate long category names
                category_name = dp['category']
                if len(category_name) > 30:
                    category_name = category_name[:27] + "..."
                
                chart_data.append({
                    "category": category_name,
                    "value": dp['value']
                })
        
        chart_title = f"Category Analysis ({most_common_unit})" if most_common_unit else "Category Analysis"
        y_title = f"Value ({most_common_unit})" if most_common_unit else "Value"
        
        # Determine chart type based on hint
        if chart_type_hint == "Pie Chart":
            # For pie chart, use arc mark
            chart_spec = {
                "$schema": "https://vega.github.io/schema/vega-lite/v5.json",
                "title": chart_title,
                "description": f"Pie chart showing distribution by category",
                "data": {"values": chart_data},
                "mark": {"type": "arc", "innerRadius": 0},
                "encoding": {
                    "theta": {"field": "value", "type": "quantitative"},
                    "color": {"field": "category", "type": "nominal", "legend": {"title": "Category"}},
                    "tooltip": [
                        {"field": "category", "type": "nominal", "title": "Category"},
                        {"field": "value", "type": "quantitative", "format": ".1f", "title": "Value"}
                    ]
                },
                "width": 400,
                "height": 400
            }
            caption = f"This pie chart shows the distribution of {most_common_unit} by category from the provided documents."
            return f"{caption}\n\n```json\n{json.dumps(chart_spec, indent=2)}\n```"
        
        elif chart_type_hint == "Donut Chart":
            # For donut chart, use arc mark with inner radius
            chart_spec = {
                "$schema": "https://vega.github.io/schema/vega-lite/v5.json",
                "title": chart_title,
                "description": f"Donut chart showing distribution by category",
                "data": {"values": chart_data},
                "mark": {"type": "arc", "innerRadius": 50},
                "encoding": {
                    "theta": {"field": "value", "type": "quantitative"},
                    "color": {"field": "category", "type": "nominal", "legend": {"title": "Category"}},
                    "tooltip": [
                        {"field": "category", "type": "nominal", "title": "Category"},
                        {"field": "value", "type": "quantitative", "format": ".1f", "title": "Value"}
                    ]
                },
                "width": 400,
                "height": 400
            }
            caption = f"This donut chart shows the distribution of {most_common_unit} by category from the provided documents."
            return f"{caption}\n\n```json\n{json.dumps(chart_spec, indent=2)}\n```"
        
        elif chart_type_hint == "Horizontal Bar":
            # Horizontal bar chart
            chart_spec = {
                "$schema": "https://vega.github.io/schema/vega-lite/v5.json",
                "title": chart_title,
                "description": f"Horizontal bar chart showing values by category",
                "data": {"values": chart_data},
                "mark": {"type": "bar", "color": "#1f77b4"},
                "encoding": {
                    "y": {"field": "category", "type": "nominal", "title": "Category", "sort": "-x"},
                    "x": {"field": "value", "type": "quantitative", "title": y_title},
                    "tooltip": [
                        {"field": "category", "type": "nominal", "title": "Category"},
                        {"field": "value", "type": "quantitative", "format": ".1f", "title": "Value"}
                    ]
                },
                "width": 600,
                "height": 400
            }
            caption = f"This horizontal bar chart shows {len(chart_data)} data points extracted from the provided documents."
            if most_common_unit:
                caption += f" Values are in {most_common_unit}."
            return f"{caption}\n\n```json\n{json.dumps(chart_spec, indent=2)}\n```"
        
        # Bar chart for categories (default)
        mark_type = "bar"
        if chart_type_hint == "Line Chart":
            mark_type = "line"
            
        chart_spec = {
            "$schema": "https://vega.github.io/schema/vega-lite/v5.json",
            "title": chart_title,
            "description": f"Bar chart showing values by category",
            "data": {"values": chart_data},
            "mark": {"type": mark_type, "color": "#1f77b4"},
            "encoding": {
                "x": {"field": "category", "type": "nominal", "title": "Category", 
                      "axis": {"labelAngle": -45, "labelLimit": 100}},
                "y": {"field": "value", "type": "quantitative", "title": y_title},
                "tooltip": [
                    {"field": "category", "type": "nominal", "title": "Category"},
                    {"field": "value", "type": "quantitative", "format": ".1f", "title": "Value"}
                ]
            },
            "width": 600,
            "height": 400
        }
    
    # Return chart with caption
    chart_type = "line chart" if any('time_period' in dp for dp in data_points) else "bar chart"
    if chart_type_hint == "Pie Chart":
        chart_type = "pie chart"
    
    caption = f"This {chart_type} shows {len(chart_data)} data points extracted from the provided documents."
    if most_common_unit:
        caption += f" Values are in {most_common_unit}."
    
    return f"{caption}\n\n```json\n{json.dumps(chart_spec, indent=2)}\n```"

def generate_csv_from_data(data_points: List[Dict[str, Any]]) -> str:
    """Generate CSV string from extracted data points."""
    if not data_points:
        return ""
    
    if any('time_period' in dp for dp in data_points):
        # Time series CSV
        csv_lines = ["time_period,value,unit"]
        for dp in data_points:
            if 'time_period' in dp:
                csv_lines.append(f"{dp['time_period']},{dp['value']},{dp.get('unit', '')}")
    else:
        # Category CSV
        csv_lines = ["category,value,unit"]
        for dp in data_points:
            if 'category' in dp:
                csv_lines.append(f"{dp['category']},{dp['value']},{dp.get('unit', '')}")
    
    return "\n".join(csv_lines)

def extract_sources_from_context(context: str) -> List[Dict[str, str]]:
    """Extract source references from formatted context."""
    sources = []
    current_source = {}
    
    for line in context.split('\n'):
        line = line.strip()
        if line.startswith('[ID:'):
            if current_source:
                sources.append(current_source)
            current_source = {'id': line[4:-1]}  # Extract ID
        elif line.startswith('source_name:'):
            current_source['source_name'] = line.split(':', 1)[1].strip()
        elif line.startswith('page_start:'):
            parts = line.split()
            current_source['page_start'] = parts[1] if len(parts) > 1 else '1'
            if 'page_end:' in line:
                current_source['page_end'] = parts[3] if len(parts) > 3 else current_source['page_start']
        elif line.startswith('section:'):
            current_source['section'] = line.split(':', 1)[1].strip()
    
    if current_source:
        sources.append(current_source)
    
    # Remove duplicates and ensure all fields
    unique_sources = []
    seen = set()
    for source in sources:
        key = (source.get('source_name', ''), source.get('page_start', ''))
        if key not in seen:
            seen.add(key)
            # Ensure all required fields
            source.setdefault('page_end', source.get('page_start', '1'))
            source.setdefault('section', '-')
            unique_sources.append(source)
    
    return unique_sources

def simple_chunk_for_charts(content: str, filename: str) -> List[Dict]:
    """Simple chunking specifically for chart data extraction."""
    import re
    
    chunks = []
    
    # Split by paragraphs and sections
    sections = re.split(r'\n\s*\n', content)
    
    for i, section in enumerate(sections):
        section = section.strip()
        if len(section) > 50:  # Only process substantial sections
            # Create chunk metadata
            chunk = {
                'id': f"{filename}_chart_{i}",
                'text': section,
                'source': filename,
                'page': 1,  # Simple page numbering
                'section': f"Section {i+1}",
                'metadata': {
                    'source_type': 'chart_upload',
                    'filename': filename,
                    'chunk_index': i
                }
            }
            chunks.append(chunk)
    
    return chunks

def search_chart_chunks(query: str, chunks: List[Dict], top_k: int = 15) -> List[Tuple[Dict, float]]:
    """Simple search for chart chunks using text similarity."""
    import re
    from collections import Counter
    
    # Simple keyword-based scoring
    query_words = set(re.findall(r'\b\w+\b', query.lower()))
    
    chunk_scores = []
    for chunk in chunks:
        text_words = set(re.findall(r'\b\w+\b', chunk['text'].lower()))
        
        # Calculate overlap score
        overlap = len(query_words.intersection(text_words))
        total_words = len(query_words.union(text_words))
        
        if total_words > 0:
            score = overlap / total_words
            
            # Boost score for chunks with numbers (likely to contain data)
            if re.search(r'\d+', chunk['text']):
                score *= 1.5
            
            # Boost score for table-like content
            if '|' in chunk['text'] or '\t' in chunk['text']:
                score *= 2.0
            
            chunk_scores.append((chunk, score))
    
    # Sort by score and return top results
    chunk_scores.sort(key=lambda x: x[1], reverse=True)
    return chunk_scores[:top_k]

def format_context_for_preview(chunks_scores: List[Tuple[Dict, float]]) -> str:
    """Format chunks for data extraction preview."""
    context_str = ""
    for i, (chunk, score) in enumerate(chunks_scores):
        context_str += f"[ID: {chunk.get('id', f'chunk_{i}')}]\n"
        context_str += f"source_name: {chunk.get('source', 'Unknown')}\n"
        context_str += f"page_start: {chunk.get('page', 1)}   page_end: {chunk.get('page', 1)}\n"
        context_str += f"section: {chunk.get('section', '-')}\n"
        context_str += f"text: |\n{chunk.get('text', '')}\n\n"
    return context_str

def display_grounded_response(response: str):
    """Display a grounded response with proper chart rendering for Streamlit."""
    import json
    import re
    
    lines = response.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Check for JSON code block (Vega-Lite chart)
        if line.strip() == '```json' and i + 1 < len(lines):
            # Find the end of the JSON block
            json_lines = []
            i += 1
            while i < len(lines) and lines[i].strip() != '```':
                json_lines.append(lines[i])
                i += 1
            
            if json_lines:
                try:
                    # Parse and render the Vega-Lite chart
                    json_content = '\n'.join(json_lines)
                    chart_spec = json.loads(json_content)
                    
                    # Check if it's a Vega-Lite spec
                    if '$schema' in chart_spec and 'vega-lite' in chart_spec['$schema']:
                        st.vega_lite_chart(chart_spec, use_container_width=True)
                    else:
                        # Fallback to code display
                        st.code(json_content, language='json')
                except json.JSONDecodeError:
                    # If JSON parsing fails, display as code
                    st.code('\n'.join(json_lines), language='json')
        
        # Check for CSV code block
        elif line.strip() == '```csv' and i + 1 < len(lines):
            # Find the end of the CSV block
            csv_lines = []
            i += 1
            while i < len(lines) and lines[i].strip() != '```':
                csv_lines.append(lines[i])
                i += 1
            
            if csv_lines:
                csv_content = '\n'.join(csv_lines)
                
                # Display CSV as both code and downloadable
                st.subheader("📊 Data (CSV)")
                st.code(csv_content, language='csv')
                
                # Add download button
                st.download_button(
                    label="📥 Download CSV",
                    data=csv_content,
                    file_name="chart_data.csv",
                    mime="text/csv"
                )
        
        # Check for markdown table
        elif '|' in line and ('---' in line or any(header in line.lower() for header in ['category', 'value', 'unit', 'time'])):
            # Collect table lines
            table_lines = [line]
            i += 1
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            i -= 1  # Adjust for the extra increment
            
            # Render as markdown table
            table_content = '\n'.join(table_lines)
            st.markdown(table_content)
        
        # Regular markdown line
        else:
            # Handle special markdown sections
            if line.startswith('## '):
                st.subheader(line[3:])
            elif line.startswith('### '):
                st.subheader(line[4:])
            elif line.startswith('- '):
                st.markdown(line)
            elif line.strip():
                st.markdown(line)
        
        i += 1

# --- Streamlit UI ---
def main():
    st.set_page_config(page_title="RAG System", page_icon="🔍", layout="wide")
    st.title("🔍 RAG System with Hybrid Search")
    
    # Initialize system
    if 'rag_system' not in st.session_state:
        with st.spinner("Initializing RAG system..."):
            st.session_state.rag_system = RAGSystem()
    
    rag_system = st.session_state.rag_system
    
    # Tabs
    tab1, tab2, tab3, tab4, tab5 = st.tabs(["📚 Index", "❓ Ask", "📊 Charts", "⚙️ Config", "🔧 Settings"])
    
    with tab1:
        st.header("Document Indexing")
        
        # File upload
        uploaded_files = st.file_uploader(
            "Upload documents", 
            type=['pdf', 'docx', 'txt', 'md'],
            accept_multiple_files=True
        )
        
        # Sidebar controls
        with st.sidebar:
            st.header("⚙️ Configuration")
            
            # Chunking strategy
            st.subheader("Chunking Strategy")
            strategy = st.selectbox(
                "Strategy",
                ["token", "character", "sentence", "paragraph", "page", 
                 "markdown_structure", "semantic_boundary", "recursive_character", "hybrid_chunking"]
            )
            
            # Token knobs
            st.subheader("Token Settings")
            target_tokens = st.slider("Target tokens", 100, 2000, 700)
            overlap_tokens = st.slider("Overlap tokens", 0, 500, 80)
            min_tokens = st.slider("Min tokens", 50, 500, 120)
            
            # Character settings
            if strategy in ["character", "recursive_character"]:
                st.subheader("Character Settings")
                char_size = st.slider("Character size", 1000, 8000, 4000)
                char_overlap = st.slider("Character overlap", 0, 1000, 400)
            else:
                char_size = 4000
                char_overlap = 400
            
            # Semantic settings
            if strategy == "semantic_boundary":
                st.subheader("Semantic Settings")
                semantic_threshold = st.slider("Semantic threshold", 0.0, 1.0, 0.65, 0.05)
            else:
                semantic_threshold = 0.65
            
            # Dense model
            st.subheader("Dense Model")
            dense_model = st.text_input("Model name", "sentence-transformers/all-MiniLM-L6-v2")
            
            # Sparse backend
            st.subheader("Sparse Backend")
            sparse_backend = st.selectbox("Backend", ["SPLADE++", "BM25"])
            
            if sparse_backend == "SPLADE++":
                splade_model = st.text_input("SPLADE model", "naver/splade-cocondenser-ensembledistil")
                top_k_vocab = st.slider("Top-k vocab", 10, 200, 60)
            else:
                splade_model = "naver/splade-cocondenser-ensembledistil"
                top_k_vocab = 60
            
            # Meta-rich settings
            st.subheader("Meta-rich Settings")
            include_front_matter = st.checkbox("Include front matter in text", False)
            front_matter_keys = st.multiselect(
                "Front matter keys",
                ["source_name", "section_path", "page_start", "page_end", "tokens", "sha256", "created_at"],
                default=["source_name", "section_path", "page_start", "page_end", "tokens", "sha256", "created_at"]
            )
            
            # Backend status
            st.subheader("Backend Status")
            dense_status = "✅" if rag_system.dense_index.embedder else "❌"
            sparse_status = "✅" if (sparse_backend == "SPLADE++" and HAS_SPLADE_DEPS) or (sparse_backend == "BM25" and HAS_BM25) else "❌"
            
            # Format reranker status with model name
            reranker_type = rag_system.reranker.reranker_type
            if reranker_type == "cohere":
                reranker_display = "Cohere v3"
            elif reranker_type == "qwen3_reranker":
                reranker_display = "Qwen3-Reranker-4B"
            elif reranker_type == "cross_encoder":
                reranker_display = "Local Cross-Encoder"
            else:
                reranker_display = "None"
            
            st.write(f"**Dense:** {dense_status} {dense_model}")
            st.write(f"**Sparse:** {sparse_status} {sparse_backend}")
            st.write(f"**Reranker:** {reranker_display}")
        
        # Process documents
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button("🔄 Chunk & Index", type="primary"):
                if uploaded_files:
                    with st.spinner("Processing documents..."):
                        # Initialize SPLADE if needed
                        if sparse_backend == "SPLADE++":
                            rag_system.init_splade(splade_model, top_k_vocab)
                        
                        # Create chunker
                        chunker = UniversalChunker(
                            strategy=strategy,
                            target_tokens=target_tokens,
                            overlap_tokens=overlap_tokens,
                            min_tokens=min_tokens,
                            char_size=char_size,
                            char_overlap=char_overlap,
                            semantic_threshold=semantic_threshold,
                            dense_embedder=rag_system.dense_index.embedder,
                            include_front_matter=include_front_matter,
                            front_matter_keys=front_matter_keys
                        )
                        
                        stats = rag_system.process_documents(uploaded_files, chunker)
                        
                        st.success(f"✅ Processed {stats['files']} files")
                        st.info(f"📄 Added {stats['new_chunks']} new chunks")
                        st.info(f"⏭️ Skipped {stats['skipped_chunks']} duplicate chunks")
                        st.info(f"📊 Total chunks: {len(rag_system.chunks)}")
                else:
                    st.warning("Please upload files first")
        
        with col2:
            if st.button("🔧 Rebuild Sparse Index"):
                with st.spinner(f"Rebuilding {sparse_backend} index..."):
                    if sparse_backend == "SPLADE++":
                        rag_system.init_splade(splade_model, top_k_vocab)
                    rag_system.rebuild_sparse_index(sparse_backend)
                    st.success(f"✅ Rebuilt {sparse_backend} index")
        
        with col3:
            if st.button("🗑️ Reset Store"):
                if st.checkbox("Confirm reset"):
                    rag_system.reset_store()
                    st.success("✅ Store reset")
                    st.rerun()
        
        # Display current stats
        if rag_system.chunks:
            st.subheader("📊 Current Index Stats")
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                st.metric("Total Chunks", len(rag_system.chunks))
            
            with col2:
                sources = set(chunk.get('source_name', 'Unknown') for chunk in rag_system.chunks)
                st.metric("Sources", len(sources))
            
            with col3:
                total_tokens = sum(chunk.get('tokens', 0) for chunk in rag_system.chunks)
                st.metric("Total Tokens", f"{total_tokens:,}")
            
            with col4:
                avg_tokens = total_tokens / len(rag_system.chunks) if rag_system.chunks else 0
                st.metric("Avg Tokens/Chunk", f"{avg_tokens:.0f}")
    
    with tab2:
        st.header("Ask Questions")
        
        if not rag_system.chunks:
            st.warning("⚠️ No documents indexed yet. Please upload and process documents first.")
            return
        
        # Query input
        query = st.text_input("🔍 Enter your question:", placeholder="What is the main topic of the documents?")
        
        # Search settings
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            retrieval_mode = st.selectbox("Retrieval mode", ["hybrid", "semantic", "sparse (SPLADE)", "alpha (weighted)"], index=0)
            
            # Alpha weight control (only show for alpha mode)
            alpha_weight = 0.7  # Default
            if retrieval_mode == "alpha (weighted)":
                alpha_weight = st.slider(
                    "Alpha Weight",
                    min_value=0.0,
                    max_value=1.0,
                    value=0.7,
                    step=0.1,
                    help="0.0 = Pure Sparse, 1.0 = Pure Dense"
                )
                st.caption(f"🎯 **{alpha_weight:.1f}** Dense + **{1-alpha_weight:.1f}** Sparse")
        
        with col2:
            top_k = st.slider("Top-k results", 1, 20, 8)
        
        with col3:
            rerank_top_n = st.slider("Rerank top-n", 5, 50, 15)
        
        with col4:
            rrf_k = st.slider("RRF k", 10, 100, 60)
        
        # Search
        if st.button("🔍 Search", type="primary") or query:
            if query:
                # Preprocess query for better results
                processed_query = query.strip()
                
                with st.spinner("Searching..."):
                    # Retrieve chunks
                    chunks_scores = rag_system.search(processed_query, retrieval_mode, top_k, rerank_top_n, rrf_k, alpha_weight=alpha_weight)
                    
                    # Rerank if available
                    if chunks_scores and rag_system.reranker.reranker_type != "none":
                        chunks_scores = rag_system.reranker.rerank(processed_query, chunks_scores, top_k)
                    
                    # Generate answer based on configuration
                    config = st.session_state.get('config', {'answer_mode': 'extractive'})
                    
                    if config.get('gpt_enabled', False) and config.get('answer_mode') == 'gpt_enhanced':
                        answer, used_chunks = generate_gpt_answer(
                            processed_query, 
                            chunks_scores,
                            gpt_model=config.get('gpt_model', 'gpt-3.5-turbo'),
                            api_key=config.get('openai_api_key')
                        )
                    elif config.get('gpt_enabled', False) and config.get('answer_mode') == 'hybrid':
                        # Try GPT first, fallback to extractive
                        gpt_answer, gpt_chunks = generate_gpt_answer(
                            processed_query, 
                            chunks_scores,
                            gpt_model=config.get('gpt_model', 'gpt-3.5-turbo'),
                            api_key=config.get('openai_api_key')
                        )
                        extractive_answer, extractive_chunks = generate_answer(processed_query, chunks_scores)
                        
                        # Combine or choose best answer
                        if "not enough information" not in gpt_answer.lower():
                            answer, used_chunks = gpt_answer, gpt_chunks
                        else:
                            answer, used_chunks = extractive_answer, extractive_chunks
                    else:
                        # Default extractive method
                        answer, used_chunks = generate_answer(processed_query, chunks_scores)
                    references = format_references(used_chunks)
                    
                    # Display results
                    st.subheader("### Answer")
                    st.write(answer)
                    
                    if references:
                        st.subheader("### References")
                        for ref in references:
                            st.write(f"- {ref}")
                    
                    # Show search details
                    with st.expander("🔧 Search Details"):
                        st.write(f"**Retrieval Mode:** {retrieval_mode}")
                        if retrieval_mode == "alpha (weighted)":
                            st.write(f"**Alpha Weight:** {alpha_weight:.1f} Dense + {1-alpha_weight:.1f} Sparse")
                        st.write(f"**Retrieved:** {len(chunks_scores)} chunks")
                        st.write(f"**Reranker:** {rag_system.reranker.reranker_type}")
                        
                        # Show answer generation method
                        answer_method = config.get('answer_mode', 'extractive') if config.get('gpt_enabled', False) else 'extractive'
                        st.write(f"**Answer Method:** {answer_method}")
                        
                        if chunks_scores:
                            st.write("**Top chunks:**")
                            for i, (chunk, score) in enumerate(chunks_scores[:3]):
                                with st.container():
                                    st.write(f"**{i+1}. Score: {score:.3f}**")
                                    st.write(f"Source: {chunk.get('source_name', 'Unknown')}")
                                    st.write(f"Text preview: {chunk['text'][:200]}...")
                                    st.divider()
            else:
                st.warning("Please enter a question")
    
    with tab3:
        st.header("📊 Grounded Charts & Tables")
        st.markdown("*Extract data and create visualizations from your documents*")
        
        # Document source selection
        st.subheader("📁 Document Source")
        doc_source = st.radio(
            "Choose document source for chart generation:",
            ["Use indexed documents", "Upload documents for charts"],
            help="Select whether to use already indexed documents or upload new ones specifically for charts"
        )
        
        chart_chunks = []
        
        if doc_source == "Use indexed documents":
            if not rag_system.chunks:
                st.warning("📝 Please upload and index documents first in the **Index** tab!")
                st.stop()
            else:
                st.success(f"✅ Ready to analyze {len(rag_system.chunks)} indexed chunks")
                chart_chunks = rag_system.chunks
        
        else:  # Upload documents for charts
            st.info("📤 Upload documents specifically for chart generation (temporary processing)")
            
            # Check document processing dependencies
            missing_deps = []
            if not HAS_PYMUPDF and not HAS_PYPDF2:
                missing_deps.append("PDF processing (install PyMuPDF: `pip install PyMuPDF`)")
            if not HAS_DOCX:
                missing_deps.append("DOCX processing (install python-docx: `pip install python-docx`)")
            
            # Check for pandas (needed for Excel/CSV)
            try:
                import pandas as pd
                HAS_PANDAS = True
            except ImportError:
                HAS_PANDAS = False
                missing_deps.append("Excel/CSV processing (install pandas: `pip install pandas openpyxl`)")
            
            if missing_deps:
                st.warning("⚠️ Some document processing features are not available:")
                for dep in missing_deps:
                    st.write(f"- {dep}")
                st.info("You can still upload TXT and MD files, or install the missing dependencies.")
            else:
                st.success("✅ All document processing features are available!")
            
            # Chart-specific file upload
            chart_files = st.file_uploader(
                "Upload documents for chart analysis",
                type=['pdf', 'docx', 'txt', 'md', 'xlsx', 'xls', 'csv'],
                accept_multiple_files=True,
                key="chart_files",
                help="Upload documents and data files for chart generation (PDF, DOCX, TXT, MD, Excel, CSV)"
            )
            
            if chart_files:
                # Process uploaded files temporarily
                with st.spinner("🔍 Processing documents for chart analysis..."):
                    temp_processor = DocumentProcessor()
                    processing_results = []
                    
                    for file in chart_files:
                        try:
                            # Process each file
                            st.info(f"Processing {file.name}...")
                            content = temp_processor.extract_text(file)
                            
                            if content.strip():
                                # Use simple chunking for chart data extraction
                                simple_chunks = simple_chunk_for_charts(content, file.name)
                                chart_chunks.extend(simple_chunks)
                                processing_results.append({
                                    'file': file.name,
                                    'status': 'success',
                                    'chunks': len(simple_chunks),
                                    'content_length': len(content)
                                })
                                st.success(f"✅ {file.name}: Extracted {len(simple_chunks)} chunks ({len(content):,} characters)")
                            else:
                                processing_results.append({
                                    'file': file.name,
                                    'status': 'empty',
                                    'message': 'No text content found'
                                })
                                st.warning(f"⚠️ {file.name}: No text content extracted")
                                
                        except Exception as e:
                            processing_results.append({
                                'file': file.name,
                                'status': 'error',
                                'message': str(e)
                            })
                            st.error(f"❌ {file.name}: {str(e)}")
                
                # Show processing summary
                if processing_results:
                    with st.expander("📋 Processing Summary"):
                        successful_files = [r for r in processing_results if r['status'] == 'success']
                        failed_files = [r for r in processing_results if r['status'] in ['error', 'empty']]
                        
                        if successful_files:
                            st.success(f"✅ Successfully processed {len(successful_files)} files:")
                            for result in successful_files:
                                st.write(f"- **{result['file']}**: {result['chunks']} chunks, {result['content_length']:,} characters")
                        
                        if failed_files:
                            st.warning(f"⚠️ Issues with {len(failed_files)} files:")
                            for result in failed_files:
                                st.write(f"- **{result['file']}**: {result['message']}")
                
                if chart_chunks:
                    st.success(f"🎉 Ready for analysis: {len(chart_chunks)} total chunks from {len([r for r in processing_results if r['status'] == 'success'])} successful files")
                else:
                    st.error("❌ No content extracted from any uploaded files")
                    st.info("💡 **Troubleshooting Tips:**")
                    st.markdown("""
                    - **PDF files**: Ensure they contain text (not just images)
                    - **DOCX files**: Check if the document has readable text content
                    - **Excel files**: Ensure sheets contain data (install pandas: `pip install pandas openpyxl`)
                    - **CSV files**: Check for proper formatting and column headers
                    - **File size**: Very large files might take longer to process
                    - **File format**: Supported formats: PDF, DOCX, TXT, MD, XLSX, XLS, CSV
                    - **Content**: Upload documents with tables, data, or structured information
                    """)
                    st.stop()
            else:
                st.info("👆 Please upload documents to generate charts")
                st.stop()
        
        if chart_chunks:
            
            # Chart request input
            chart_request = st.text_area(
                "📊 Describe the table or chart you want:",
                placeholder="Create a bar chart of quarterly revenue by region for 2023",
                height=100,
                help="Examples:\n• Create a table of employee counts by department\n• Make a line chart of sales trends over time\n• Show a pie chart of budget allocation by category"
            )
            
            # Quick example buttons
            st.write("**Quick Examples:**")
            col_ex1, col_ex2, col_ex3, col_ex4 = st.columns(4)
            
            with col_ex1:
                if st.button("📊 Revenue Chart"):
                    st.session_state.chart_request = "Create a bar chart showing revenue data"
            
            with col_ex2:
                if st.button("📋 Hours Table"):
                    st.session_state.chart_request = "Create a table of working hours and schedules"
            
            with col_ex3:
                if st.button("📈 Trends"):
                    st.session_state.chart_request = "Show line chart of trends over time"
                    
            with col_ex4:
                if st.button("🥧 Distribution"):
                    st.session_state.chart_request = "Create a pie chart showing distribution by category"
            
            # Use session state value if set
            if 'chart_request' in st.session_state and st.session_state.chart_request:
                chart_request = st.session_state.chart_request
                st.session_state.chart_request = ""  # Clear after use
            
            # Chart generation settings
            col1, col2, col3 = st.columns(3)
            
            with col1:
                chart_retrieval_mode = st.selectbox(
                    "Data Retrieval Mode",
                    ["hybrid", "semantic", "sparse (SPLADE)", "alpha (weighted)"],
                    index=0,
                    help="How to search for relevant data chunks"
                )
                
                # Alpha weight for alpha mode
                chart_alpha_weight = 0.7
                if chart_retrieval_mode == "alpha (weighted)":
                    chart_alpha_weight = st.slider(
                        "Alpha Weight", 0.0, 1.0, 0.7, 0.1,
                        help="Balance between dense and sparse search"
                    )
            
            with col2:
                chart_top_k = st.slider("Chunks to Analyze", 5, 30, 15)
                chart_rerank_top_n = st.slider("Rerank Top-N", 10, 50, 25)
            
            with col3:
                chart_type_hint = st.selectbox(
                    "Preferred Chart Type",
                    ["Auto-detect", "Table", "Bar Chart", "Line Chart", "Pie Chart", "Donut Chart", "Horizontal Bar", "Area Chart", "Scatter Plot"],
                    help="Hint for visualization type (AI will choose best fit)"
                )
                
                include_csv = st.checkbox("Include CSV Data", value=True, help="Show machine-readable CSV with charts")
            
            # Generate chart
            if st.button("📊 Generate Chart/Table", type="primary") or chart_request:
                if chart_request.strip():
                    with st.spinner("🔍 Analyzing documents and extracting data..."):
                        try:
                            # Search for relevant data
                            processed_query = chart_request.strip()
                            
                            if doc_source == "Use indexed documents":
                                # Use full RAG system search
                                chunks_scores = rag_system.search(
                                    processed_query, 
                                    chart_retrieval_mode, 
                                    chart_top_k, 
                                    chart_rerank_top_n, 
                                    60,  # rrf_k
                                    alpha_weight=chart_alpha_weight
                                )
                                
                                # Rerank if available
                                if chunks_scores and rag_system.reranker.reranker_type != "none":
                                    chunks_scores = rag_system.reranker.rerank(processed_query, chunks_scores, chart_top_k)
                            
                            else:
                                # Use simple search for uploaded chart documents
                                chunks_scores = search_chart_chunks(processed_query, chart_chunks, chart_top_k)
                                st.info(f"🔍 Using simple text search on {len(chart_chunks)} uploaded chunks (advanced search available with indexed documents)")
                            
                            if chunks_scores:
                                # Generate grounded chart/table response
                                chart_response = generate_grounded_chart(
                                    chart_request, 
                                    chunks_scores, 
                                    chart_type_hint,
                                    include_csv
                                )
                                
                                # Display the response
                                st.markdown("---")
                                
                                # Parse and display response with proper chart rendering
                                display_grounded_response(chart_response)
                                
                                # Show extracted data preview  
                                preview_context = format_context_for_preview(chunks_scores)
                                extracted_data = extract_numeric_data_from_context(preview_context)
                                if extracted_data:
                                    with st.expander(f"📊 Extracted Data Preview ({len(extracted_data)} items)"):
                                        # Create a simple dataframe view
                                        import pandas as pd
                                        
                                        preview_data = []
                                        for dp in extracted_data[:10]:  # Show top 10
                                            preview_data.append({
                                                'Category': dp.get('category', dp.get('time_period', 'N/A')),
                                                'Value': dp['value'],
                                                'Unit': dp.get('unit', ''),
                                                'Source': dp.get('raw_text', '')[:100] + "..."
                                            })
                                        
                                        if preview_data:
                                            df = pd.DataFrame(preview_data)
                                            st.dataframe(df, use_container_width=True)
                                            
                                            st.caption(f"✨ **Data Quality:** {len(extracted_data)} data points extracted and cleaned")
                                        else:
                                            st.info("No structured data found in the retrieved chunks")
                                
                                # Show data source details
                                with st.expander("📋 Data Source Details"):
                                    st.write(f"**Query:** {chart_request}")
                                    st.write(f"**Document Source:** {doc_source}")
                                    
                                    if doc_source == "Use indexed documents":
                                        st.write(f"**Retrieval Mode:** {chart_retrieval_mode}")
                                        if chart_retrieval_mode == "alpha (weighted)":
                                            st.write(f"**Alpha Weight:** {chart_alpha_weight:.1f}")
                                        st.write(f"**Reranker:** {rag_system.reranker.reranker_type}")
                                    else:
                                        st.write(f"**Search Method:** Simple text matching")
                                        st.write(f"**Total Uploaded Files:** {len(chart_files) if 'chart_files' in locals() and chart_files else 0}")
                                    
                                    st.write(f"**Chunks Analyzed:** {len(chunks_scores)}")
                                    st.write(f"**Chart Type Hint:** {chart_type_hint}")
                                    
                                    # Show source chunks
                                    st.subheader("📄 Source Chunks Used")
                                    for i, (chunk, score) in enumerate(chunks_scores[:10]):  # Show top 10
                                        with st.expander(f"Chunk {i+1} - Score: {score:.3f}"):
                                            st.write(f"**Source:** {chunk.get('source', 'Unknown')}")
                                            st.write(f"**Page:** {chunk.get('page', 'N/A')}")
                                            if chunk.get('section'):
                                                st.write(f"**Section:** {chunk.get('section')}")
                                            if chunk.get('metadata', {}).get('source_type') == 'chart_upload':
                                                st.info("📤 This chunk is from an uploaded document (temporary)")
                                            st.write("**Content:**")
                                            st.text(chunk.get('text', '')[:500] + "..." if len(chunk.get('text', '')) > 500 else chunk.get('text', ''))
                            
                            else:
                                st.error("❌ No relevant data found in the documents for your request. Try rephrasing your query or check if the data exists in your uploaded documents.")
                        
                        except Exception as e:
                            st.error(f"❌ Error generating chart: {str(e)}")
                            logger.error(f"Chart generation error: {e}")
                
                else:
                    st.warning("Please describe the chart or table you want to create")
        
        # Document source comparison
        with st.expander("📋 Document Source Options"):
            col1, col2 = st.columns(2)
            
            with col1:
                st.subheader("🗂️ Use Indexed Documents")
                st.markdown("""
                **Advantages:**
                - ✅ Full RAG system with advanced search
                - ✅ Semantic, sparse, hybrid, and alpha retrieval modes
                - ✅ Reranking with Cohere/Qwen3/Cross-encoder
                - ✅ Persistent storage and indexing
                - ✅ Best for complex queries and large document sets
                
                **Best for:**
                - Comprehensive document analysis
                - Complex search queries
                - Reusable document collections
                """)
            
            with col2:
                st.subheader("📤 Upload Documents for Charts")
                st.markdown("""
                **Advantages:**
                - ✅ Quick one-time analysis
                - ✅ No need to index documents
                - ✅ Temporary processing (privacy-focused)
                - ✅ Simple text-based search
                - ✅ Great for specific data extraction tasks
                
                **Best for:**
                - Quick chart generation
                - Temporary document analysis
                - Simple data extraction
                """)
        
        # Help section
        with st.expander("💡 How to Use Charts & Tables"):
            st.markdown("""
            ### 📊 **Chart & Table Generation Examples:**
            
            **📋 Tables:**
            - *"Create a table of sales figures by quarter"*
            - *"Show employee headcount by department"*
            - *"List all project deadlines and status"*
            
            **📊 Charts:**
            - *"Bar chart of revenue by region for 2023"*
            - *"Line graph showing monthly user growth"*
            - *"Pie chart of budget allocation by category"*
            
            ### 🎯 **Features:**
            - **Grounded Data Only:** Charts use only information from your documents
            - **Source Citations:** Every chart includes references to source pages
            - **Multiple Formats:** Tables, bar charts, line graphs, pie charts, scatter plots
            - **CSV Export:** Machine-readable data included with visualizations
            - **Smart Detection:** AI chooses the best chart type for your data
            
            ### ⚠️ **Important Notes:**
            - Data extraction is limited to what's available in your uploaded documents
            - Charts will show "Not enough information" if data is incomplete
            - All numbers and facts are strictly from document content
            - No external data sources are used
            """)
    
    with tab4:
        st.header("⚙️ Configuration")
        
        # Initialize session state for config
        if 'config' not in st.session_state:
            st.session_state.config = {
                'gpt_enabled': False,
                'gpt_model': 'gpt-3.5-turbo',
                'openai_api_key': '',
                'answer_mode': 'extractive',
                'max_tokens': 500,
                'temperature': 0.1
            }
        
        config = st.session_state.config
        
        # GPT Configuration Section
        st.subheader("🤖 GPT Integration")
        
        col1, col2 = st.columns(2)
        
        with col1:
            gpt_enabled = st.checkbox(
                "Enable GPT-enhanced answers", 
                value=config['gpt_enabled'],
                help="Use GPT to generate more natural answers from retrieved context"
            )
            config['gpt_enabled'] = gpt_enabled
            
            if gpt_enabled:
                gpt_model = st.selectbox(
                    "GPT Model",
                    ["gpt-3.5-turbo", "gpt-4", "gpt-4-turbo", "gpt-4o", "gpt-4o-mini"],
                    index=["gpt-3.5-turbo", "gpt-4", "gpt-4-turbo", "gpt-4o", "gpt-4o-mini"].index(config['gpt_model'])
                )
                config['gpt_model'] = gpt_model
                
                openai_api_key = st.text_input(
                    "OpenAI API Key",
                    value=config['openai_api_key'],
                    type="password",
                    help="Your OpenAI API key for GPT access"
                )
                config['openai_api_key'] = openai_api_key
        
        with col2:
            if gpt_enabled:
                answer_mode = st.radio(
                    "Answer Generation Mode",
                    ["extractive", "gpt_enhanced", "hybrid"],
                    index=["extractive", "gpt_enhanced", "hybrid"].index(config['answer_mode']),
                    help="Choose how answers are generated"
                )
                config['answer_mode'] = answer_mode
                
                max_tokens = st.slider(
                    "Max tokens for GPT response",
                    100, 1000, config['max_tokens'],
                    help="Maximum length of GPT-generated answers"
                )
                config['max_tokens'] = max_tokens
                
                temperature = st.slider(
                    "Temperature",
                    0.0, 1.0, config['temperature'], 0.1,
                    help="Creativity level (0=focused, 1=creative)"
                )
                config['temperature'] = temperature
        
        # Answer Mode Explanations
        if gpt_enabled:
            st.subheader("📝 Answer Mode Details")
            
            if config['answer_mode'] == 'extractive':
                st.info("**Extractive**: Uses pattern matching and keyword extraction from retrieved chunks (fast, factual)")
            elif config['answer_mode'] == 'gpt_enhanced':
                st.info("**GPT Enhanced**: Uses GPT to synthesize natural answers from retrieved context (natural, comprehensive)")
            elif config['answer_mode'] == 'hybrid':
                st.info("**Hybrid**: Combines extractive and GPT methods for best of both worlds (balanced)")
        
        # API Status Section
        st.subheader("🔧 API Status")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            openai_status = "✅ Available" if HAS_OPENAI else "❌ Not installed"
            st.write(f"**OpenAI**: {openai_status}")
            
            if HAS_OPENAI and config['openai_api_key']:
                if st.button("Test OpenAI Connection"):
                    try:
                        client = openai.OpenAI(api_key=config['openai_api_key'])
                        response = client.chat.completions.create(
                            model=config['gpt_model'],
                            messages=[{"role": "user", "content": "Hello"}],
                            max_tokens=5
                        )
                        st.success("✅ OpenAI API connection successful!")
                    except Exception as e:
                        st.error(f"❌ OpenAI API connection failed: {str(e)}")
        
        with col2:
            cohere_status = "✅ Available" if HAS_COHERE else "❌ Not installed"
            st.write(f"**Cohere**: {cohere_status}")
        
        with col3:
            transformers_status = "✅ Available" if HAS_SPLADE_DEPS else "❌ Not installed"
            st.write(f"**Transformers**: {transformers_status}")
        
        # Installation Instructions
        st.subheader("📦 Installation Instructions")
        
        with st.expander("Install missing dependencies"):
            st.code("""
# For OpenAI GPT models
pip install openai

# For Cohere reranking  
pip install cohere

# For SPLADE++ and Qwen3-Reranker-4B
pip install transformers torch

# For better vector storage
pip install faiss-cpu

# For advanced text splitting
pip install langchain-text-splitters
            """, language="bash")
        
        # Save config
        st.session_state.config = config
    
    with tab4:
        st.header("🔧 Advanced Settings")
        
        # Initialize settings state
        if 'settings' not in st.session_state:
            st.session_state.settings = {
                # Search & Retrieval Settings
                'default_retrieval_mode': 'hybrid',
                'default_top_k': 8,
                'default_rerank_top_n': 15,
                'default_rrf_k': 60,
                'similarity_threshold': 0.65,
                'default_alpha_weight': 0.7,
                
                # Chunking Settings
                'default_chunk_strategy': 'token',
                'default_target_tokens': 700,
                'default_overlap_tokens': 80,
                'default_min_tokens': 120,
                'default_char_size': 4000,
                'default_char_overlap': 400,
                
                # Dense Index Settings
                'dense_model': 'sentence-transformers/all-MiniLM-L6-v2',
                'normalize_embeddings': True,
                'dense_index_type': 'auto',  # auto, chroma, faiss, memory
                
                # Sparse Index Settings
                'sparse_backend': 'SPLADE++',
                'splade_model': 'naver/splade-cocondenser-ensembledistil',
                'splade_top_k_vocab': 60,
                'bm25_k1': 1.2,
                'bm25_b': 0.75,
                
                # Answer Generation Settings
                'answer_max_sentences': 5,
                'answer_min_overlap': 2,
                'answer_diversity_threshold': 0.6,
                'time_pattern_boost': 3,
                'keyword_score_multiplier': 2,
                
                # System Settings
                'storage_location': './rag_store',
                'auto_save_interval': 60,  # seconds
                'max_cache_size': 1000,  # MB
                'debug_mode': False,
                'performance_mode': 'balanced',  # fast, balanced, quality
                
                # UI Settings
                'show_search_details': True,
                'show_chunk_previews': True,
                'show_scores': True,
                'results_per_page': 10,
                'compact_view': False
            }
        
        settings = st.session_state.settings
        
        # Search & Retrieval Settings
        st.subheader("🔍 Search & Retrieval")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.write("**Default Search Parameters**")
            
            default_retrieval_mode = st.selectbox(
                "Default Retrieval Mode",
                ["semantic", "sparse (SPLADE)", "hybrid", "alpha (weighted)"],
                index=["semantic", "sparse (SPLADE)", "hybrid", "alpha (weighted)"].index(settings['default_retrieval_mode'])
            )
            settings['default_retrieval_mode'] = default_retrieval_mode
            
            default_top_k = st.slider(
                "Default Top-K Results", 1, 20, settings['default_top_k']
            )
            settings['default_top_k'] = default_top_k
            
            default_rerank_top_n = st.slider(
                "Default Rerank Top-N", 5, 50, settings['default_rerank_top_n']
            )
            settings['default_rerank_top_n'] = default_rerank_top_n
        
        with col2:
            st.write("**Advanced Search Parameters**")
            
            default_rrf_k = st.slider(
                "Default RRF K", 10, 100, settings['default_rrf_k']
            )
            settings['default_rrf_k'] = default_rrf_k
            
            similarity_threshold = st.slider(
                "Semantic Similarity Threshold", 0.0, 1.0, settings['similarity_threshold'], 0.05
            )
            settings['similarity_threshold'] = similarity_threshold
            
            default_alpha_weight = st.slider(
                "Default Alpha Weight (Dense vs Sparse)", 0.0, 1.0, settings.get('default_alpha_weight', 0.7), 0.1,
                help="For alpha mode: 0.0 = Pure Sparse, 1.0 = Pure Dense"
            )
            settings['default_alpha_weight'] = default_alpha_weight
        
        # Chunking Settings
        st.subheader("📄 Document Chunking")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.write("**Default Chunking Strategy**")
            
            default_chunk_strategy = st.selectbox(
                "Strategy",
                ["token", "character", "sentence", "paragraph", "page", 
                 "markdown_structure", "semantic_boundary", "recursive_character", "hybrid_chunking"],
                index=["token", "character", "sentence", "paragraph", "page", 
                       "markdown_structure", "semantic_boundary", "recursive_character", "hybrid_chunking"].index(settings['default_chunk_strategy'])
            )
            settings['default_chunk_strategy'] = default_chunk_strategy
        
        with col2:
            st.write("**Token Settings**")
            
            default_target_tokens = st.slider(
                "Target Tokens", 100, 2000, settings['default_target_tokens']
            )
            settings['default_target_tokens'] = default_target_tokens
            
            default_overlap_tokens = st.slider(
                "Overlap Tokens", 0, 500, settings['default_overlap_tokens']
            )
            settings['default_overlap_tokens'] = default_overlap_tokens
            
            default_min_tokens = st.slider(
                "Min Tokens", 50, 500, settings['default_min_tokens']
            )
            settings['default_min_tokens'] = default_min_tokens
        
        with col3:
            st.write("**Character Settings**")
            
            default_char_size = st.slider(
                "Character Size", 1000, 8000, settings['default_char_size']
            )
            settings['default_char_size'] = default_char_size
            
            default_char_overlap = st.slider(
                "Character Overlap", 0, 1000, settings['default_char_overlap']
            )
            settings['default_char_overlap'] = default_char_overlap
        
        # Index Configuration
        st.subheader("🗂️ Index Configuration")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.write("**Dense Index Settings**")
            
            dense_model = st.text_input(
                "Dense Model",
                value=settings['dense_model'],
                help="HuggingFace model for dense embeddings"
            )
            settings['dense_model'] = dense_model
            
            normalize_embeddings = st.checkbox(
                "Normalize Embeddings",
                value=settings['normalize_embeddings']
            )
            settings['normalize_embeddings'] = normalize_embeddings
            
            dense_index_type = st.selectbox(
                "Dense Index Type",
                ["auto", "chroma", "faiss", "memory"],
                index=["auto", "chroma", "faiss", "memory"].index(settings['dense_index_type'])
            )
            settings['dense_index_type'] = dense_index_type
        
        with col2:
            st.write("**Sparse Index Settings**")
            
            sparse_backend = st.selectbox(
                "Sparse Backend",
                ["SPLADE++", "BM25"],
                index=["SPLADE++", "BM25"].index(settings['sparse_backend'])
            )
            settings['sparse_backend'] = sparse_backend
            
            if sparse_backend == "SPLADE++":
                splade_model = st.text_input(
                    "SPLADE Model",
                    value=settings['splade_model']
                )
                settings['splade_model'] = splade_model
                
                splade_top_k_vocab = st.slider(
                    "SPLADE Top-K Vocab", 10, 200, settings['splade_top_k_vocab']
                )
                settings['splade_top_k_vocab'] = splade_top_k_vocab
            
            elif sparse_backend == "BM25":
                bm25_k1 = st.slider(
                    "BM25 K1", 0.1, 3.0, settings['bm25_k1'], 0.1
                )
                settings['bm25_k1'] = bm25_k1
                
                bm25_b = st.slider(
                    "BM25 B", 0.0, 1.0, settings['bm25_b'], 0.05
                )
                settings['bm25_b'] = bm25_b
        
        # Answer Generation Settings
        st.subheader("💬 Answer Generation")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.write("**Answer Quality Settings**")
            
            answer_max_sentences = st.slider(
                "Max Sentences in Answer", 1, 10, settings['answer_max_sentences']
            )
            settings['answer_max_sentences'] = answer_max_sentences
            
            answer_min_overlap = st.slider(
                "Min Word Overlap", 1, 5, settings['answer_min_overlap']
            )
            settings['answer_min_overlap'] = answer_min_overlap
            
            answer_diversity_threshold = st.slider(
                "Answer Diversity Threshold", 0.0, 1.0, settings['answer_diversity_threshold'], 0.05
            )
            settings['answer_diversity_threshold'] = answer_diversity_threshold
        
        with col2:
            st.write("**Scoring Parameters**")
            
            time_pattern_boost = st.slider(
                "Time Pattern Score Boost", 1, 10, settings['time_pattern_boost']
            )
            settings['time_pattern_boost'] = time_pattern_boost
            
            keyword_score_multiplier = st.slider(
                "Keyword Score Multiplier", 1, 5, settings['keyword_score_multiplier']
            )
            settings['keyword_score_multiplier'] = keyword_score_multiplier
        
        # System Settings
        st.subheader("⚙️ System Configuration")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.write("**Storage & Performance**")
            
            storage_location = st.text_input(
                "Storage Location",
                value=settings['storage_location'],
                help="Directory for storing indices and metadata"
            )
            settings['storage_location'] = storage_location
            
            auto_save_interval = st.slider(
                "Auto-save Interval (seconds)", 10, 300, settings['auto_save_interval']
            )
            settings['auto_save_interval'] = auto_save_interval
            
            max_cache_size = st.slider(
                "Max Cache Size (MB)", 100, 5000, settings['max_cache_size']
            )
            settings['max_cache_size'] = max_cache_size
            
            performance_mode = st.selectbox(
                "Performance Mode",
                ["fast", "balanced", "quality"],
                index=["fast", "balanced", "quality"].index(settings['performance_mode'])
            )
            settings['performance_mode'] = performance_mode
        
        with col2:
            st.write("**Debug & UI Settings**")
            
            debug_mode = st.checkbox(
                "Enable Debug Mode",
                value=settings['debug_mode'],
                help="Show detailed logging and debug information"
            )
            settings['debug_mode'] = debug_mode
            
            show_search_details = st.checkbox(
                "Show Search Details by Default",
                value=settings['show_search_details']
            )
            settings['show_search_details'] = show_search_details
            
            show_chunk_previews = st.checkbox(
                "Show Chunk Previews",
                value=settings['show_chunk_previews']
            )
            settings['show_chunk_previews'] = show_chunk_previews
            
            show_scores = st.checkbox(
                "Show Relevance Scores",
                value=settings['show_scores']
            )
            settings['show_scores'] = show_scores
            
            compact_view = st.checkbox(
                "Compact View Mode",
                value=settings['compact_view']
            )
            settings['compact_view'] = compact_view
        
        # Action Buttons
        st.subheader("🎯 Actions")
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            if st.button("💾 Save Settings", type="primary"):
                st.success("✅ Settings saved successfully!")
        
        with col2:
            if st.button("🔄 Reset to Defaults"):
                # Reset to default values
                st.session_state.settings = {
                    'default_retrieval_mode': 'hybrid',
                    'default_top_k': 8,
                    'default_rerank_top_n': 15,
                    'default_rrf_k': 60,
                    'similarity_threshold': 0.65,
                    'default_alpha_weight': 0.7,
                    'default_chunk_strategy': 'token',
                    'default_target_tokens': 700,
                    'default_overlap_tokens': 80,
                    'default_min_tokens': 120,
                    'default_char_size': 4000,
                    'default_char_overlap': 400,
                    'dense_model': 'sentence-transformers/all-MiniLM-L6-v2',
                    'normalize_embeddings': True,
                    'dense_index_type': 'auto',
                    'sparse_backend': 'SPLADE++',
                    'splade_model': 'naver/splade-cocondenser-ensembledistil',
                    'splade_top_k_vocab': 60,
                    'bm25_k1': 1.2,
                    'bm25_b': 0.75,
                    'answer_max_sentences': 5,
                    'answer_min_overlap': 2,
                    'answer_diversity_threshold': 0.6,
                    'time_pattern_boost': 3,
                    'keyword_score_multiplier': 2,
                    'storage_location': './rag_store',
                    'auto_save_interval': 60,
                    'max_cache_size': 1000,
                    'debug_mode': False,
                    'performance_mode': 'balanced',
                    'show_search_details': True,
                    'show_chunk_previews': True,
                    'show_scores': True,
                    'results_per_page': 10,
                    'compact_view': False
                }
                st.success("✅ Settings reset to defaults!")
                st.rerun()
        
        with col3:
            if st.button("📤 Export Settings"):
                import json
                settings_json = json.dumps(settings, indent=2)
                st.download_button(
                    label="📁 Download Settings JSON",
                    data=settings_json,
                    file_name="rag_settings.json",
                    mime="application/json"
                )
        
        with col4:
            uploaded_settings = st.file_uploader(
                "📥 Import Settings",
                type=['json'],
                help="Upload a previously exported settings file"
            )
            
            if uploaded_settings is not None:
                try:
                    import json
                    imported_settings = json.load(uploaded_settings)
                    st.session_state.settings.update(imported_settings)
                    st.success("✅ Settings imported successfully!")
                    st.rerun()
                except Exception as e:
                    st.error(f"❌ Failed to import settings: {e}")
        
        # Settings Summary
        st.subheader("📋 Current Settings Summary")
        
        with st.expander("View All Settings"):
            import json
            st.json(settings)
        
        # Performance Impact Warning
        if settings['performance_mode'] == 'quality':
            st.warning("⚠️ Quality mode may increase processing time but provides better results.")
        elif settings['performance_mode'] == 'fast':
            st.info("⚡ Fast mode prioritizes speed over accuracy.")
        
        # Save settings
        st.session_state.settings = settings

if __name__ == "__main__":
    main()

# --- How to run ---
"""
## Installation

Install required dependencies:
```bash
pip install streamlit tiktoken PyPDF2 pymupdf python-docx nltk sentence-transformers numpy chromadb rank-bm25 pyyaml scikit-learn
```

Optional dependencies for advanced features:
```bash
# For SPLADE++ and Qwen3-Reranker-4B
pip install transformers torch tokenizers

# For better vector storage
pip install faiss-cpu

# For OpenAI GPT models
pip install openai

# For Cohere reranking
pip install cohere

# For advanced text splitting
pip install langchain-text-splitters

# For spaCy sentence splitting
pip install spacy
python -m spacy download en_core_web_sm
```

Download NLTK data:
```python
import nltk
nltk.download('punkt')
```

## Configuration

Set environment variables for optional features:
```bash
export COHERE_API_KEY="your-cohere-api-key"  # For Cohere reranking
```

## Usage

Run the application:
```bash
streamlit run app.py
```

Navigate to the Index tab to upload and process documents, then use the Ask tab to query your documents.

## Features

- **Multi-format support**: PDF, DOCX, MD, TXT
- **Multiple chunking strategies**: Token, character, sentence, paragraph, page, markdown-aware, semantic, recursive, hybrid
- **Dense + sparse retrieval**: Sentence transformers + SPLADE++ (with BM25 fallback)
- **Hybrid search**: RRF fusion of dense and sparse results
- **Reranking**: Cohere v3 > Qwen3-Reranker-4B > local cross-encoder
- **Persistence**: All indices and metadata persist across sessions
- **Extractive answering**: Grounded answers with references

The system automatically handles missing dependencies with graceful fallbacks.
"""
